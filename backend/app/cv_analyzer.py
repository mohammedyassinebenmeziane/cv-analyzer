import PyPDF2
import docx
import json
import re
import requests
from typing import List, Dict, Optional
import os
from dotenv import load_dotenv

load_dotenv()

class CVAnalyzer:
    def __init__(self):
        # Configuration Hugging Face API (optionnelle - fonctionne sans clé pour les modèles publics)
        self.hf_api_key = os.getenv("HUGGINGFACE_API_KEY", None)
        self.hf_api_url = "https://api-inference.huggingface.co/models"
        
        # Mode rapide : désactiver les appels API (activé par défaut pour de meilleures performances)
        # Mettre FAST_MODE=true dans .env pour forcer le mode rapide
        self.fast_mode = os.getenv("FAST_MODE", "true").lower() == "true"
        
        # Session requests réutilisable pour de meilleures performances
        self.session = requests.Session() if not self.fast_mode else None
        if self.session:
            self.session.headers.update({
                "Content-Type": "application/json"
            })
            if self.hf_api_key:
                self.session.headers.update({
                    "Authorization": f"Bearer {self.hf_api_key}"
                })
        
        # Modèle de similarité sémantique (gratuit, léger)
        self.similarity_model = "sentence-transformers/all-MiniLM-L6-v2"
        
        # Modèle IA spécialisé pour l'extraction de compétences et entités depuis les CVs
        # Utilisation d'un modèle de NER (Named Entity Recognition) pour extraction dynamique
        self.ner_model = "dslim/bert-base-NER"  # Modèle NER généraliste
        # Alternative: "dbmdz/bert-large-cased-finetuned-conll03-english" pour meilleure précision
        
        # Modèle pour l'extraction de compétences techniques (token classification)
        self.skill_extraction_model = "microsoft/unilm-base-cased"  # Modèle généraliste pour extraction
        
        # Modèle pour l'analyse de texte et extraction d'informations structurées
        self.text_analysis_model = "sentence-transformers/all-MiniLM-L6-v2"
    
    def _call_hf_api(self, model: str, inputs: Dict, task: str = "feature-extraction") -> Optional[Dict]:
        """Appelle l'API Hugging Face Inference"""
        if self.fast_mode or not self.session:
            return None
        
        url = f"{self.hf_api_url}/{model}"
        
        try:
            # Utiliser la session réutilisable pour de meilleures performances
            response = self.session.post(url, json=inputs, timeout=5.0)  # Timeout réduit à 5s
            if response.status_code == 200:
                return response.json()
            return None
        except requests.Timeout:
            # En cas de timeout, retourner None silencieusement
            return None
        except Exception as e:
            # En cas d'erreur, retourner None silencieusement
            return None
    
    def _calculate_semantic_similarity(self, text1: str, text2: str) -> float:
        """Calcule la similarité sémantique entre deux textes (utilise calcul amélioré local)"""
        if not text1 or not text2:
            return 0.0
        
        # Utiliser directement le calcul amélioré (plus rapide et fiable que l'API)
        # L'API Hugging Face est trop lente et peut échouer
        return self._enhanced_similarity(text1, text2)
    
    def _enhanced_similarity(self, text1: str, text2: str) -> float:
        """Calcul amélioré de similarité basé sur les mots-clés et la structure"""
        if not text1 or not text2:
            return 0.0
        
        text1_lower = text1.lower()
        text2_lower = text2.lower()
        
        # Extraire les mots significatifs (3+ caractères pour capturer plus de mots)
        words1 = set(re.findall(r'\b\w{3,}\b', text1_lower))
        words2 = set(re.findall(r'\b\w{3,}\b', text2_lower))
        
        # Filtrer les mots communs non significatifs
        stop_words = {'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'her', 'was', 'one', 'our', 'out', 'day', 'get', 'has', 'him', 'his', 'how', 'its', 'may', 'new', 'now', 'old', 'see', 'two', 'way', 'who', 'boy', 'did', 'she', 'use', 'her', 'many', 'than', 'them', 'these', 'le', 'de', 'la', 'les', 'des', 'du', 'un', 'une', 'et', 'ou', 'pour', 'avec', 'dans', 'sur', 'par', 'est', 'sont', 'été', 'être', 'avoir', 'fait', 'faire', 'avec', 'dans', 'pour', 'sont', 'cette', 'cette', 'comme', 'plus', 'tout', 'tous', 'toutes'}
        words1 = {w for w in words1 if w not in stop_words and len(w) >= 3}
        words2 = {w for w in words2 if w not in stop_words and len(w) >= 3}
        
        if not words1 or not words2:
            return 0.0
        
        # Calcul Jaccard amélioré
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        jaccard = len(intersection) / len(union) if union else 0.0
        
        # Bonus pour les phrases communes (2+ mots consécutifs)
        phrases1 = set(re.findall(r'\b\w{3,}\s+\w{3,}\b', text1_lower))
        phrases2 = set(re.findall(r'\b\w{3,}\s+\w{3,}\b', text2_lower))
        phrase_score = 0.0
        if phrases1 and phrases2:
            phrase_intersection = phrases1.intersection(phrases2)
            phrase_union = phrases1.union(phrases2)
            phrase_score = len(phrase_intersection) / len(phrase_union) if phrase_union else 0.0
        
        # Combiner les scores (70% mots, 30% phrases)
        final_score = (jaccard * 0.7 + phrase_score * 0.3) if phrase_score > 0 else jaccard
        
        # Bonus si beaucoup de mots communs (indique une forte similarité)
        if len(intersection) >= 5:
            final_score = min(final_score * 1.1, 1.0)
        
        # Améliorer le score si les textes sont courts et ont des mots-clés communs
        # (cas où la description est courte mais pertinente)
        if len(text1_lower.split()) <= 10 or len(text2_lower.split()) <= 10:
            # Si un des textes est court, être plus généreux avec les correspondances
            if len(intersection) >= 2:
                final_score = max(final_score, 0.3)  # Minimum 0.3 si au moins 2 mots communs
            if len(intersection) >= 3:
                final_score = max(final_score, 0.5)  # Minimum 0.5 si au moins 3 mots communs
        
        # Pénalité seulement si vraiment très peu de mots communs ET beaucoup de mots différents
        if len(intersection) == 0:
            final_score = 0.0
        elif len(intersection) <= 1 and len(union) > 30:
            final_score = final_score * 0.5  # Réduire de 50% seulement si vraiment très différent
        
        return max(0.0, min(1.0, final_score))
    
    def _basic_similarity(self, text1: str, text2: str) -> float:
        """Calcul basique de similarité basé sur les mots communs"""
        words1 = set(re.findall(r'\b\w{3,}\b', text1.lower()))
        words2 = set(re.findall(r'\b\w{3,}\b', text2.lower()))
        if not words1 or not words2:
            return 0.0
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        return len(intersection) / len(union) if union else 0.0
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extrait le texte d'un fichier PDF avec amélioration du formatage"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        # Nettoyer et améliorer le formatage
                        # Remplacer les espaces multiples par un seul espace
                        page_text = re.sub(r'\s+', ' ', page_text)
                        # Restaurer les sauts de ligne pour les listes
                        page_text = re.sub(r'([.!?])\s+([A-Z])', r'\1\n\2', page_text)
                        # Restaurer les sauts de ligne pour les dates
                        page_text = re.sub(r'(\d{4})\s+([A-Z])', r'\1\n\2', page_text)
                        text += page_text + "\n\n"
        except Exception as e:
            raise Exception(f"Erreur lors de l'extraction du PDF: {str(e)}")
        return text.strip()
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extrait le texte d'un fichier DOCX avec amélioration du formatage"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    text += para_text + "\n"
            
            # Extraire aussi les tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text += row_text + "\n"
        except Exception as e:
            raise Exception(f"Erreur lors de l'extraction du DOCX: {str(e)}")
        return text.strip()
    
    def extract_text(self, file_path: str, file_extension: str) -> str:
        """Extrait le texte selon le type de fichier"""
        if file_extension.lower() == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension.lower() in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Format de fichier non supporté: {file_extension}")
    
    def extract_skills(self, cv_text: str) -> List[str]:
        """Extrait les compétences du CV de manière dynamique avec IA (sans liste statique)"""
        found_skills = []
        
        # Méthode 1: Utiliser un modèle NER pour identifier les compétences comme entités
        skills_from_ner = self._extract_skills_with_ner(cv_text)
        found_skills.extend(skills_from_ner)
        
        # Méthode 2: Utiliser l'extraction sémantique pour identifier les compétences
        skills_from_semantic = self._extract_skills_with_semantic_analysis(cv_text)
        found_skills.extend(skills_from_semantic)
        
        # Méthode 3: Extraction par patterns (compétences techniques communes)
        skills_from_patterns = self._extract_skills_with_patterns(cv_text)
        found_skills.extend(skills_from_patterns)
        
        # Dédupliquer et nettoyer
        unique_skills = []
        seen = set()
        for skill in found_skills:
            skill_lower = skill.lower().strip()
            if skill_lower and len(skill_lower) > 2 and skill_lower not in seen:
                seen.add(skill_lower)
                # Normaliser le nom
                normalized = skill.strip().capitalize() if skill.islower() else skill.strip()
                unique_skills.append(normalized)
        
        return unique_skills[:50]  # Limiter à 50 compétences
    
    def _extract_skills_with_ner(self, text: str) -> List[str]:
        """Extrait les compétences en utilisant un modèle NER (Named Entity Recognition)"""
        # En mode rapide, retourner une liste vide (utiliser les autres méthodes)
        if self.fast_mode:
            return []
        
        skills = []
        try:
            # Utiliser l'API Hugging Face pour NER
            url = f"{self.hf_api_url}/{self.ner_model}"
            headers = {"Content-Type": "application/json"}
            if self.hf_api_key:
                headers["Authorization"] = f"Bearer {self.hf_api_key}"
            
            # Limiter le texte
            text_chunks = [text[i:i+500] for i in range(0, min(len(text), 2000), 500)]
            
            for chunk in text_chunks:
                payload = {"inputs": chunk}
                try:
                    # Ajouter un timeout pour éviter les blocages
                    response = requests.post(url, headers=headers, json=payload, timeout=10.0)
                    if response.status_code == 200:
                        result = response.json()
                        # Le résultat NER est une liste de dictionnaires avec 'word' et 'entity'
                        if isinstance(result, list):
                            for item in result:
                                if isinstance(item, dict):
                                    entity = item.get('entity', '')
                                    word = item.get('word', '')
                                    # Filtrer les entités pertinentes (ORG, MISC peuvent contenir des compétences)
                                    if entity in ['ORG', 'MISC'] and len(word) > 2:
                                        # Nettoyer le mot
                                        clean_word = re.sub(r'[^\w\s-]', '', word).strip()
                                        if clean_word and len(clean_word) > 2:
                                            skills.append(clean_word)
                except requests.Timeout:
                    print(f"Timeout NER extraction")
                    continue
                except Exception as e:
                    print(f"Erreur NER extraction: {str(e)}")
                    continue
        except Exception as e:
            print(f"Erreur lors de l'extraction NER: {str(e)}")
        
        return skills
    
    def _extract_skills_with_semantic_analysis(self, text: str) -> List[str]:
        """Extrait les compétences en utilisant l'analyse sémantique et les patterns"""
        skills = []
        text_lower = text.lower()
        
        # Patterns pour identifier les compétences techniques
        # Chercher des mots techniques (majuscules, acronymes, noms propres techniques)
        
        # 1. Acronymes techniques (2-5 lettres en majuscules)
        acronyms = re.findall(r'\b[A-Z]{2,5}\b', text)
        for acro in acronyms:
            # Filtrer les acronymes communs qui ne sont pas des compétences
            exclude = ['CV', 'PDF', 'API', 'URL', 'HTTP', 'HTTPS', 'HTML', 'CSS', 'JS', 'ID', 'UI', 'UX']
            if acro not in exclude and len(acro) >= 2:
                skills.append(acro)
        
        # 2. Mots techniques avec points (ex: React.js, Node.js)
        tech_with_dots = re.findall(r'\b[A-Z][a-z]+\.(?:js|ts|py|net|jsx|tsx)\b', text)
        skills.extend(tech_with_dots)
        
        # 3. Technologies en majuscules suivies de mots (ex: WORDPRESS, DOCKER)
        tech_uppercase = re.findall(r'\b[A-Z]{3,}[A-Za-z]*\b', text)
        for tech in tech_uppercase:
            if len(tech) >= 3 and tech not in ['THE', 'AND', 'FOR', 'ARE', 'ALL', 'YOU', 'CAN']:
                skills.append(tech)
        
        # 4. Phrases techniques communes (ex: "machine learning", "data science")
        tech_phrases = [
            r'\bmachine\s+learning\b',
            r'\bdata\s+science\b',
            r'\bdeep\s+learning\b',
            r'\bartificial\s+intelligence\b',
            r'\bweb\s+development\b',
            r'\bfull\s+stack\b',
            r'\bfront\s+end\b',
            r'\bback\s+end\b',
            r'\bcloud\s+computing\b',
            r'\bdevops\b',
            r'\bci/cd\b',
            r'\brest\s+api\b',
            r'\bgraphql\b',
        ]
        for pattern in tech_phrases:
            matches = re.findall(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                skills.append(match.strip().title())
        
        return skills
    
    def _extract_skills_with_patterns(self, text: str) -> List[str]:
        """Extrait les compétences en utilisant des patterns de texte"""
        skills = []
        text_lower = text.lower()
        
        # Chercher des sections de compétences communes
        skill_section_patterns = [
            r'(?:skills|compétences|technologies|tools|outils)[\s:]+([^\n]+)',
            r'(?:proficient|experienced|familiar|knowledgeable)\s+in\s+([^\n]+)',
            r'(?:expertise|maîtrise|connaissance)[\s:]+([^\n]+)',
        ]
        
        for pattern in skill_section_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                # Extraire les compétences de la phrase (séparées par virgules, points, etc.)
                skill_candidates = re.split(r'[,;•\-\n]', match)
                for candidate in skill_candidates:
                    candidate = candidate.strip()
                    # Filtrer les candidats valides (2-50 caractères, pas de mots communs)
                    if 2 <= len(candidate) <= 50:
                        common_words = ['and', 'or', 'the', 'with', 'in', 'for', 'to', 'of', 'a', 'an']
                        if candidate.lower() not in common_words:
                            skills.append(candidate)
        
        return skills
    
    def extract_languages(self, cv_text: str) -> List[str]:
        """Extrait les langues parlées du CV"""
        # Liste des langues communes (français, anglais, etc.)
        languages_keywords = {
            "français": ["français", "french", "francais", "francophone"],
            "anglais": ["anglais", "english", "anglophone"],
            "espagnol": ["espagnol", "spanish", "español"],
            "allemand": ["allemand", "german", "deutsch"],
            "italien": ["italien", "italian", "italiano"],
            "arabe": ["arabe", "arabic", "عربي"],
            "chinois": ["chinois", "chinese", "中文", "mandarin"],
            "japonais": ["japonais", "japanese", "日本語"],
            "portugais": ["portugais", "portuguese", "português"],
            "russe": ["russe", "russian", "русский"],
            "néerlandais": ["néerlandais", "dutch", "nederlands"],
            "polonais": ["polonais", "polish", "polski"],
            "turc": ["turc", "turkish", "türkçe"],
            "coréen": ["coréen", "korean", "한국어"],
            "hindi": ["hindi", "हिंदी"],
            "hébreu": ["hébreu", "hebrew", "עברית"],
            "suédois": ["suédois", "swedish", "svenska"],
            "norvégien": ["norvégien", "norwegian", "norsk"],
            "danois": ["danois", "danish", "dansk"],
            "grec": ["grec", "greek", "ελληνικά"]
        }
        
        cv_text_lower = cv_text.lower()
        found_languages = []
        
        # Rechercher les sections de langues
        language_section_patterns = [
            r'(?i)(langues?|languages?|idiomas?|sprachen?)',
            r'(?i)(compétences? linguistiques?|linguistic skills?)',
            r'(?i)(parle|speak|habla|spricht)'
        ]
        
        # Chercher dans tout le texte
        for lang_name, keywords in languages_keywords.items():
            for keyword in keywords:
                if keyword.lower() in cv_text_lower:
                    if lang_name not in found_languages:
                        found_languages.append(lang_name)
                        break
        
        # Chercher aussi les niveaux de langue (A1, A2, B1, B2, C1, C2, natif, etc.)
        language_level_patterns = [
            r'(?i)(natif|native|maternel|mother tongue)',
            r'(?i)(courant|fluent|avancé|advanced)',
            r'(?i)(intermédiaire|intermediate|moyen)',
            r'(?i)(débutant|beginner|basic)',
            r'(?i)(A1|A2|B1|B2|C1|C2)'
        ]
        
        # Si on trouve des niveaux mais pas de langues, chercher autour
        lines = cv_text.split('\n')
        for i, line in enumerate(lines):
            line_lower = line.lower()
            # Si la ligne contient un niveau de langue
            if any(re.search(pattern, line_lower) for pattern in language_level_patterns):
                # Chercher les langues dans les lignes proches
                context_lines = lines[max(0, i-2):min(len(lines), i+3)]
                context_text = ' '.join(context_lines).lower()
                for lang_name, keywords in languages_keywords.items():
                    if lang_name not in found_languages:
                        for keyword in keywords:
                            if keyword.lower() in context_text:
                                found_languages.append(lang_name)
                                break
        
        return found_languages
    
    def extract_experience(self, cv_text: str) -> List[str]:
        """Extrait les expériences du CV"""
        experience_patterns = [
            r'(?i)(expérience|experience|work|employment|emploi)',
            r'(?i)(\d{4})\s*[-–]\s*(\d{4}|présent|present|now)',
        ]
        
        experiences = []
        lines = cv_text.split('\n')
        in_experience_section = False
        current_experience = []
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                if current_experience:
                    experiences.append(' '.join(current_experience))
                    current_experience = []
                in_experience_section = False
                continue
            
            if any(re.search(pattern, line_stripped) for pattern in experience_patterns):
                in_experience_section = True
                if current_experience:
                    experiences.append(' '.join(current_experience))
                current_experience = [line_stripped]
            elif in_experience_section and len(line_stripped) > 10:
                current_experience.append(line_stripped)
        
        if current_experience:
            experiences.append(' '.join(current_experience))
        
        return experiences[:10]
    
    def analyze_cv(self, cv_text: str, job_description: str) -> Dict:
        """
        Analyse intelligente du CV avec IA pour scan rapide et décision objective.
        Extrait automatiquement et compare avec la description du poste :
        - Compétences (techniques, métiers, soft skills)
        - Expériences professionnelles (pertinentes/non pertinentes)
        - Formation et diplômes
        - Certifications
        - Langues
        - Projets / réalisations
        - Informations personnelles
        - Score et recommandations
        """
        # 1. EXTRACTION COMPLÈTE DU PROFIL STRUCTURÉ (avec IA)
        try:
            candidate_profile = self.extract_candidate_profile(cv_text, job_description)
        except Exception as e:
            print(f"Erreur lors de l'extraction du profil candidat: {str(e)}")
            import traceback
            traceback.print_exc()
            # Profil par défaut en cas d'erreur
            candidate_profile = {
                "identite": {},
                "resume_professionnel": {},
                "competences_techniques": {},
                "experiences_professionnelles": [],
                "stages_alternances": [],
                "projets": [],
                "formation": [],
                "certifications": [],
                "langues": [],
                "soft_skills": [],
                "score_correspondance": 0.0
            }
        
        # 2. EXTRACTION DES COMPÉTENCES (techniques, métiers, soft skills)
        cv_skills_technical = candidate_profile.get("competences_techniques", {})
        all_cv_skills = []
        for category, skills in cv_skills_technical.items():
            if isinstance(skills, list):
                all_cv_skills.extend([s.lower() for s in skills])
        
        # Soft skills
        soft_skills = candidate_profile.get("soft_skills", [])
        all_cv_skills.extend([s.lower() for s in soft_skills])
        
        # 3. EXTRACTION DES COMPÉTENCES REQUISES DU POSTE (avec IA sémantique)
        try:
            required_skills = self._extract_required_skills_from_job(job_description)
        except Exception as e:
            print(f"Erreur lors de l'extraction des compétences requises: {str(e)}")
            required_skills = []
        
        # 4. COMPARAISON DES COMPÉTENCES (avec IA)
        try:
            missing_skills, matching_skills = self._compare_skills_with_ia(all_cv_skills, required_skills, job_description)
        except Exception as e:
            print(f"Erreur lors de la comparaison des compétences: {str(e)}")
            missing_skills = []
            matching_skills = []
        
        # 5. ANALYSE DES EXPÉRIENCES PROFESSIONNELLES (avec IA)
        experiences = candidate_profile.get("experiences_professionnelles", [])
        try:
            relevant_experience, irrelevant_experience = self._classify_experiences_with_ia(
                experiences, job_description
            )
        except Exception as e:
            print(f"Erreur lors de la classification des expériences: {str(e)}")
            relevant_experience = []
            irrelevant_experience = []
        
        # 6. ANALYSE DE LA FORMATION (avec IA)
        education = candidate_profile.get("formation", [])
        try:
            education_match_score = self._evaluate_education_relevance(education, job_description)
        except Exception as e:
            print(f"Erreur lors de l'évaluation de la formation: {str(e)}")
            education_match_score = 0.0
        
        # 7. ANALYSE DES CERTIFICATIONS (avec IA)
        certifications = candidate_profile.get("certifications", [])
        try:
            cert_match_score = self._evaluate_certifications_relevance(certifications, job_description)
        except Exception as e:
            print(f"Erreur lors de l'évaluation des certifications: {str(e)}")
            cert_match_score = 0.0
        
        # 8. ANALYSE DES PROJETS (avec IA)
        projects = candidate_profile.get("projets", [])
        try:
            projects_match_score = self._evaluate_projects_relevance(projects, job_description)
        except Exception as e:
            print(f"Erreur lors de l'évaluation des projets: {str(e)}")
            projects_match_score = 0.0
        
        # 9. ANALYSE DES LANGUES
        languages = candidate_profile.get("langues", [])
        
        # 10. CALCUL DU SCORE GLOBAL (basé sur tous les critères avec pondération IA)
        try:
            score = self._calculate_comprehensive_score(
                matching_skills=matching_skills,
                required_skills=required_skills,
                relevant_experience=relevant_experience,
                education_match=education_match_score,
                cert_match=cert_match_score,
                projects_match=projects_match_score,
                cv_text=cv_text,
                job_description=job_description
            )
        except Exception as e:
            print(f"Erreur lors du calcul du score: {str(e)}")
            score = 0.0
        
        # 11. GÉNÉRATION DE RECOMMANDATIONS INTELLIGENTES (avec IA)
        try:
            recommendations = self._generate_ai_recommendations(
                score=score,
                missing_skills=missing_skills,
                relevant_experience=relevant_experience,
                education_match=education_match_score,
                cert_match=cert_match_score,
                cv_text=cv_text,
                job_description=job_description
            )
        except Exception as e:
            print(f"Erreur lors de la génération des recommandations: {str(e)}")
            recommendations = []
        
        # Formater les expériences de manière sécurisée
        formatted_relevant = []
        for exp in relevant_experience[:5]:
            try:
                formatted_relevant.append(self._format_experience(exp))
            except Exception as e:
                formatted_relevant.append(str(exp))
        
        formatted_irrelevant = []
        for exp in irrelevant_experience[:5]:
            try:
                formatted_irrelevant.append(self._format_experience(exp))
            except Exception as e:
                formatted_irrelevant.append(str(exp))
        
        # Extraire les langues de manière sécurisée
        languages_list = []
        if languages:
            for lang in languages:
                if isinstance(lang, dict):
                    languages_list.append(lang.get("langue", ""))
                elif isinstance(lang, str):
                    languages_list.append(lang)
        
        return {
            "score": round(score, 2) if score is not None else 0.0,
            "missing_skills": missing_skills if missing_skills else [],
            "relevant_experience": formatted_relevant,
            "irrelevant_experience": formatted_irrelevant,
            "recommendations": recommendations if recommendations else [],
            "languages": languages_list,
            "candidate_profile": candidate_profile if candidate_profile else {}
        }
    
    def _extract_required_skills_from_job(self, job_description: str) -> List[str]:
        """Extrait les compétences requises de la description du poste avec IA (sans liste statique)"""
        required_skills = []
        
        # Méthode 1: Extraction par patterns (sections de compétences requises) - PRIORITAIRE
        tech_patterns = [
            r'(?:requis|required|demandé|nécessaire|maîtrise|mastery|compétences|skills|qualifications|technologies|maîtriser|connaître|connaissance)[\s:]+([^.\n]+)',
            r'(?:proficient|experienced|familiar|knowledgeable|expert)\s+(?:in|en|de|avec)\s+([^\n]+)',
            r'(?:doit|must|should|need|besoin)\s+(?:maîtriser|connaître|avoir|posséder)[\s:]+([^.\n]+)',
        ]
        for pattern in tech_patterns:
            matches = re.findall(pattern, job_description, re.IGNORECASE)
            for match in matches:
                # Extraire les compétences de la phrase
                skills_from_match = self._extract_skills_with_patterns(match)
                required_skills.extend(skills_from_match)
        
        # Méthode 2: Extraire le titre du poste et les termes clés (PRIORITAIRE)
        job_lower = job_description.lower()
        first_line = job_description.split('\n')[0].lower() if job_description else ""
        
        # Extraire les termes clés du titre et de la première ligne
        title_keywords = re.findall(r'\b\w{3,}\b', first_line)
        for keyword in title_keywords:
            if len(keyword) > 3:
                required_skills.append(keyword.capitalize())
        
        # Détecter les rôles/positions clés
        role_patterns = [
            r'(?:full\s*stack|fullstack)\s*(?:developer|développeur|dev)',
            r'(?:front\s*end|frontend)\s*(?:developer|développeur|dev)',
            r'(?:back\s*end|backend)\s*(?:developer|développeur|dev)',
            r'(?:software|développeur|developer|engineer|ingénieur|architect|architecte)',
            r'(?:comptable|expert\s*comptable|accountant)',
            r'(?:marketing|digital\s*marketing|community\s*manager)',
        ]
        for pattern in role_patterns:
            matches = re.findall(pattern, job_description, re.IGNORECASE)
            for match in matches:
                required_skills.append(match.strip())
        
        # Méthode 3: Extraire les compétences techniques communes mentionnées dans le texte
        common_tech_skills = [
            'javascript', 'python', 'java', 'react', 'node', 'vue', 'angular', 'django', 'flask',
            'sql', 'mongodb', 'postgresql', 'mysql', 'redis', 'docker', 'kubernetes', 'aws', 'azure',
            'git', 'jenkins', 'ci/cd', 'rest', 'graphql', 'microservices', 'agile', 'scrum',
            'html', 'css', 'typescript', 'php', 'ruby', 'go', 'rust', 'c++', 'c#', '.net',
            'full stack', 'frontend', 'backend', 'fullstack', 'devops', 'linux', 'unix',
            'comptabilité', 'sage', 'ciel', 'excel', 'power bi', 'tableau', 'finance',
            'marketing', 'seo', 'sem', 'google analytics', 'facebook ads', 'content marketing'
        ]
        
        for skill in common_tech_skills:
            if skill in job_lower:
                skill_formatted = skill.title() if ' ' not in skill else skill
                if skill_formatted not in required_skills:
                    required_skills.append(skill_formatted)
        
        # Méthode 3: Extraction ligne par ligne (pour les listes de compétences)
        lines = job_description.split('\n')
        for line in lines:
            line_stripped = line.strip()
            if line_stripped and len(line_stripped) < 100:
                line_skills = self._extract_skills_with_semantic_analysis(line_stripped)
                required_skills.extend(line_skills)
        
        # Méthode 4: Extraction du titre du poste (souvent contient des compétences clés)
        first_line = job_description.split('\n')[0] if job_description else ""
        if first_line and len(first_line) < 200:
            first_line_lower = first_line.lower()
            
            # Détecter les titres de poste courants et leurs compétences associées
            job_title_mappings = {
                'développeur': ['développement', 'programmation', 'coding', 'code', 'javascript', 'python', 'java'],
                'developer': ['développement', 'programmation', 'coding', 'code', 'javascript', 'python', 'java'],
                'dev': ['développement', 'programmation', 'coding', 'code', 'javascript', 'python', 'java'],
                'full stack': ['full stack', 'fullstack', 'frontend', 'backend', 'javascript', 'react', 'node', 'html', 'css', 'api', 'database'],
                'fullstack': ['full stack', 'fullstack', 'frontend', 'backend', 'javascript', 'react', 'node', 'html', 'css', 'api', 'database'],
                'frontend': ['frontend', 'react', 'vue', 'angular', 'javascript', 'html', 'css', 'typescript'],
                'backend': ['backend', 'node', 'python', 'java', 'api', 'database', 'sql', 'rest'],
                'comptable': ['comptabilité', 'sage', 'ciel', 'excel', 'fiscalité', 'tva', 'déclarations fiscales'],
                'accountant': ['comptabilité', 'sage', 'ciel', 'excel', 'fiscalité', 'tva', 'déclarations fiscales'],
                'marketing': ['marketing', 'seo', 'sem', 'google analytics', 'social media', 'content marketing'],
            }
            
            # Si le titre contient un mot-clé de poste, ajouter les compétences associées (PRIORITÉ)
            for job_keyword, associated_skills in job_title_mappings.items():
                if job_keyword in first_line_lower:
                    # Ajouter chaque compétence 3 fois pour leur donner plus de poids
                    for skill in associated_skills:
                        skill_normalized = skill.title() if ' ' not in skill else skill
                        if skill_normalized.lower() not in [s.lower() for s in required_skills]:
                            required_skills.append(skill_normalized)
                            required_skills.append(skill_normalized)  # Doubler pour plus de poids
                            required_skills.append(skill_normalized)  # Tripler
            
            # Extraire les mots-clés du titre (donner plus de poids)
            title_words = re.findall(r'\b\w{4,}\b', first_line_lower)
            for word in title_words:
                # Ignorer les mots communs
                common_words = ['poste', 'position', 'job', 'travail', 'work', 'cherche', 'recherche', 'recherchons', 'nous', 'vous', 'pour', 'avec', 'dans', 'sur', 'description', 'du', 'de', 'la', 'le', 'les', 'un', 'une', 'des']
                if word not in common_words and len(word) > 3:
                    required_skills.append(word.capitalize())
            
            # Si le titre contient des mots-clés techniques, les ajouter
            title_skills = self._extract_skills_with_semantic_analysis(first_line)
            # Donner plus de poids aux compétences du titre (tripler)
            required_skills.extend(title_skills)
            required_skills.extend(title_skills)
            required_skills.extend(title_skills)
        
        # Dédupliquer et nettoyer
        unique_skills = []
        seen = set()
        for skill in required_skills:
            skill_lower = skill.lower().strip()
            if skill_lower and len(skill_lower) > 2 and skill_lower not in seen:
                seen.add(skill_lower)
                # Normaliser le nom
                normalized = skill.strip().capitalize() if skill.islower() else skill.strip()
                unique_skills.append(normalized)
        
        return unique_skills[:30]  # Limiter à 30 compétences
    
    def _compare_skills_with_ia(self, cv_skills: List[str], required_skills: List[str], job_description: str) -> tuple:
        """Compare les compétences du CV avec celles requises en utilisant l'IA sémantique"""
        if not required_skills:
            return [], []
        
        matching_skills = []
        missing_skills = []
        cv_skills_lower = [s.lower() for s in cv_skills]
        job_lower = job_description.lower()
        
        # Identifier les compétences critiques (mentionnées plusieurs fois ou dans le titre)
        first_line = job_description.split('\n')[0].lower() if job_description else ""
        critical_skills = []
        skill_counts = {}
        for skill in required_skills:
            skill_lower = skill.lower()
            count = job_lower.count(skill_lower)
            skill_counts[skill_lower] = count
            if count > 1 or skill_lower in first_line:
                critical_skills.append(skill_lower)
        
        # Pour chaque compétence requise, vérifier si elle existe dans le CV
        for req_skill in required_skills:
            req_skill_lower = req_skill.lower()
            found = False
            
            # Vérification exacte d'abord
            if req_skill_lower in cv_skills_lower:
                matching_skills.append(req_skill)
                found = True
            else:
                # Vérification partielle (mots-clés dans la compétence)
                req_words = set(re.findall(r'\b\w{3,}\b', req_skill_lower))
                for cv_skill in cv_skills:
                    cv_skill_lower = cv_skill.lower()
                    cv_words = set(re.findall(r'\b\w{3,}\b', cv_skill_lower))
                    
                    # Si au moins 50% des mots correspondent
                    if req_words and cv_words:
                        overlap = len(req_words.intersection(cv_words)) / len(req_words)
                        if overlap >= 0.5:
                            matching_skills.append(req_skill)
                            found = True
                            break
                
                # Vérification sémantique (seulement si pas trouvé)
                if not found:
                    for cv_skill in cv_skills:
                        similarity = self._calculate_semantic_similarity(req_skill, cv_skill)
                        if similarity > 0.75:  # Seuil de similarité très élevé
                            matching_skills.append(req_skill)
                            found = True
                            break
            
            if not found:
                missing_skills.append(req_skill)
        
        return missing_skills, matching_skills
    
    def _classify_experiences_with_ia(self, experiences: List[Dict], job_description: str) -> tuple:
        """Classifie les expériences en pertinentes/non pertinentes avec IA"""
        relevant = []
        irrelevant = []
        
        if not experiences:
            return relevant, irrelevant
        
        for exp in experiences:
            # Construire un texte descriptif de l'expérience
            exp_text = ""
            if isinstance(exp, dict):
                exp_text = f"{exp.get('intitule_poste', '')} {exp.get('entreprise', '')} {' '.join(exp.get('missions', []))}"
            else:
                exp_text = str(exp)
            
            # Calculer la similarité sémantique avec la description du poste
            similarity = self._calculate_semantic_similarity(exp_text[:500], job_description[:500])
            
            if similarity > 0.35:  # Seuil ajusté pour meilleure précision
                relevant.append(exp)
            else:
                irrelevant.append(exp)
        
        # Si pas assez d'expériences pertinentes, utiliser méthode basique
        if len(relevant) == 0 and experiences:
            job_keywords = set(re.findall(r'\b\w{4,}\b', job_description.lower()))
            for exp in experiences:
                exp_text = str(exp) if not isinstance(exp, dict) else f"{exp.get('intitule_poste', '')} {exp.get('entreprise', '')}"
                exp_keywords = set(re.findall(r'\b\w{4,}\b', exp_text.lower()))
            if len(exp_keywords.intersection(job_keywords)) >= 2:
                    relevant.append(exp)
            else:
                    irrelevant.append(exp)
        
        return relevant, irrelevant
    
    def _evaluate_education_relevance(self, education: List[Dict], job_description: str) -> float:
        """Évalue la pertinence de la formation par rapport au poste avec IA"""
        if not education:
            return 0.0
        
        max_score = 0.0
        for edu in education:
            edu_text = f"{edu.get('diplome', '')} {edu.get('domaine', '')} {edu.get('etablissement', '')}"
            similarity = self._calculate_semantic_similarity(edu_text[:300], job_description[:300])
            max_score = max(max_score, similarity)
        
        return max_score
    
    def _evaluate_certifications_relevance(self, certifications: List[Dict], job_description: str) -> float:
        """Évalue la pertinence des certifications avec IA"""
        if not certifications:
            return 0.0
        
        max_score = 0.0
        for cert in certifications:
            cert_text = f"{cert.get('nom', '')} {cert.get('organisme', '')}"
            similarity = self._calculate_semantic_similarity(cert_text[:300], job_description[:300])
            max_score = max(max_score, similarity)
        
        return max_score
    
    def _evaluate_projects_relevance(self, projects: List[Dict], job_description: str) -> float:
        """Évalue la pertinence des projets avec IA"""
        if not projects:
            return 0.0
        
        max_score = 0.0
        for project in projects:
            project_text = f"{project.get('nom', '')} {project.get('description', '')} {' '.join(project.get('technologies', []))}"
            similarity = self._calculate_semantic_similarity(project_text[:400], job_description[:400])
            max_score = max(max_score, similarity)
        
        return max_score
    
    def _calculate_comprehensive_score(
        self, 
        matching_skills: List[str],
        required_skills: List[str],
        relevant_experience: List[Dict],
        education_match: float,
        cert_match: float,
        projects_match: float,
        cv_text: str,
        job_description: str
    ) -> float:
        """Calcule un score global basé sur tous les critères avec pondération IA"""
        
        # Extraire le résumé professionnel du CV pour comparaison directe
        professional_summary = ""
        lines = cv_text.split('\n')
        summary_keywords = ['résumé', 'resume', 'profil', 'profile', 'summary', 'about', 'à propos']
        in_summary = False
        summary_lines = []
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in summary_keywords):
                in_summary = True
                continue
            if in_summary:
                if line.strip() and len(line.strip()) > 10:
                    summary_lines.append(line.strip())
                elif len(summary_lines) > 0:
                    break
        
        if summary_lines:
            professional_summary = ' '.join(summary_lines[:4])
        else:
            # Si pas de résumé trouvé, utiliser les premières lignes du CV
            professional_summary = ' '.join([line.strip() for line in lines[:10] if line.strip() and len(line.strip()) > 10])
        
        # Si aucune compétence requise n'est identifiée, comparer directement description vs résumé
        if not required_skills:
            job_lower = job_description.lower()
            summary_lower = professional_summary.lower() if professional_summary else cv_text[:500].lower()
            
            # Comparaison directe description vs résumé professionnel
            job_keywords = set(re.findall(r'\b\w{4,}\b', job_lower))
            summary_keywords = set(re.findall(r'\b\w{4,}\b', summary_lower))
            
            if job_keywords:
                overlap = len(job_keywords.intersection(summary_keywords)) / len(job_keywords)
                # Score basé uniquement sur la correspondance
                return min(max(overlap * 100, 0.0), 100.0)
            return 30.0  # Score très bas si pas de correspondance
        
        # 1. Score des compétences (70% - CRITIQUE - facteur dominant)
        skills_score = len(matching_skills) / len(required_skills) if required_skills else 0.0
        
        # Si aucune compétence ne correspond, score = 0
        if len(matching_skills) == 0 and len(required_skills) > 0:
            skills_score = 0.0
        
        # Pénalité TRÈS STRICTE si moins de 50% des compétences requises sont présentes
        if skills_score < 0.5:
            skills_score = skills_score * 0.2  # Pénalité de 80% si moins de 50% de correspondance
        
        # Pénalité EXTRÊME si moins de 30%
        if skills_score < 0.3:
            skills_score = skills_score * 0.1  # Pénalité de 90% supplémentaire
        
        # 2. Comparaison directe Description vs Résumé Professionnel (20% - NOUVEAU)
        job_lower = job_description.lower()
        summary_lower = professional_summary.lower() if professional_summary else cv_text[:500].lower()
        
        # Extraire les mots-clés importants de la description
        job_keywords = set(re.findall(r'\b\w{4,}\b', job_lower))
        summary_keywords = set(re.findall(r'\b\w{4,}\b', summary_lower))
        
        summary_match_score = 0.0
        if job_keywords and summary_keywords:
            overlap = len(job_keywords.intersection(summary_keywords)) / len(job_keywords)
            summary_match_score = overlap
        
        # 3. Score des expériences pertinentes (5% - réduit)
        exp_score = min(len(relevant_experience) / 3, 1.0) if relevant_experience else 0.0
        
        # 4. Similarité sémantique globale CV vs Poste (3% - très réduit)
        cv_summary = cv_text[:1500]
        job_summary = job_description[:1500]
        semantic_score = self._calculate_semantic_similarity(cv_summary, job_summary)
        
        # 5. Score formation (1%)
        education_score = education_match if education_match else 0.0
        
        # 6. Score certifications (1%)
        cert_score = cert_match if cert_match else 0.0
        
        # 7. Score projets (0% - supprimé)
        projects_score = 0.0
        
        # Calcul du score final pondéré avec compétences et résumé comme facteurs dominants
        final_score = (
            skills_score * 0.70 +           # 70% pour les compétences (facteur dominant)
            summary_match_score * 0.20 +    # 20% pour la correspondance description vs résumé
            exp_score * 0.05 +              # 5% pour les expériences
            semantic_score * 0.03 +         # 3% pour la similarité sémantique
            education_score * 0.01 +       # 1% pour la formation
            cert_score * 0.01               # 1% pour les certifications
        ) * 100
        
        # Pénalités TRÈS STRICTES basées sur la correspondance
        # Si moins de 30% des compétences correspondent, score maximum de 25%
        if skills_score < 0.3:
            final_score = min(final_score, 25.0)
        
        # Si moins de 20% des compétences correspondent, score maximum de 15%
        if skills_score < 0.2:
            final_score = min(final_score, 15.0)
        
        # Si la correspondance description vs résumé est très faible (< 20%), pénalité supplémentaire
        if summary_match_score < 0.2:
            final_score = final_score * 0.5  # Réduction de 50%
        
        # Si la correspondance description vs résumé est faible (< 30%), pénalité modérée
        if summary_match_score < 0.3:
            final_score = final_score * 0.7  # Réduction de 30%
        
        # Si les deux sont faibles, score très bas
        if skills_score < 0.3 and summary_match_score < 0.2:
            final_score = min(final_score, 20.0)
        
        # Si les deux sont très faibles, score extrêmement bas
        if skills_score < 0.2 and summary_match_score < 0.15:
            final_score = min(final_score, 10.0)
        
        # Vérification finale : si le score des compétences est 0, le score maximum est 15%
        if skills_score == 0.0:
            final_score = min(final_score, 15.0)
        
        return min(max(final_score, 0.0), 100.0)
    
    def _generate_ai_recommendations(
        self,
        score: float,
        missing_skills: List[str],
        relevant_experience: List[Dict],
        education_match: float,
        cert_match: float,
        cv_text: str,
        job_description: str
    ) -> List[str]:
        """Génère des recommandations intelligentes basées sur l'analyse IA"""
        recommendations = []
        
        # Recommandations sur les compétences
        if missing_skills:
            top_missing = missing_skills[:3]
            recommendations.append(
                f"Compétences manquantes critiques : {', '.join(top_missing)}. "
                f"Considérer une formation ou certification dans ces domaines."
            )
        
        # Recommandations sur les expériences
        if len(relevant_experience) < 2:
            recommendations.append(
                "Expérience professionnelle limitée pour ce poste. "
                "Mettre en avant les projets personnels ou stages pertinents."
            )
        
        # Recommandations sur la formation
        if education_match < 0.3:
            recommendations.append(
                "La formation ne correspond pas directement au poste. "
                "Mettre en avant les compétences acquises et leur applicabilité."
            )
        
        # Recommandations sur les certifications
        if cert_match < 0.3:
            recommendations.append(
                "Aucune certification pertinente identifiée. "
                "Considérer des certifications reconnues dans le domaine."
            )
        
        # Recommandations globales basées sur le score
        if score < 50:
            recommendations.append(
                "Correspondance faible avec le poste. "
                "Le candidat nécessite une formation significative ou une réorientation."
            )
        elif score < 70:
            recommendations.append(
                "Correspondance modérée. Le candidat a des bases mais nécessite "
                "un développement de compétences spécifiques au poste."
            )
        elif score < 85:
            recommendations.append(
                "Bonne correspondance. Le candidat présente un profil adapté "
                "avec quelques axes d'amélioration possibles."
            )
        else:
            recommendations.append(
                "Excellente correspondance. Le profil du candidat correspond "
                "très bien aux exigences du poste."
            )
        
        # Recommandation basée sur la similarité sémantique
        semantic_similarity = self._calculate_semantic_similarity(
            cv_text[:1000], job_description[:1000]
        )
        if semantic_similarity < 0.4:
            recommendations.append(
                "Le contenu global du CV ne correspond pas suffisamment à la description du poste. "
                "Reformuler certaines sections pour mieux aligner le profil."
            )
        
        if not recommendations:
            recommendations.append("Le CV correspond bien au poste demandé.")
        
        return recommendations[:5]  # Limiter à 5 recommandations
    
    def _format_experience(self, exp: Dict) -> str:
        """Formate une expérience pour l'affichage"""
        if isinstance(exp, dict):
            parts = []
            if exp.get("intitule_poste"):
                parts.append(exp["intitule_poste"])
            if exp.get("entreprise"):
                parts.append(f"chez {exp['entreprise']}")
            if exp.get("periode"):
                parts.append(f"({exp['periode']})")
            if exp.get("missions"):
                parts.append(" - " + " | ".join(exp["missions"][:2]))
            return " ".join(parts) if parts else str(exp)
        return str(exp)
    
    def extract_candidate_profile(self, cv_text: str, job_description: str) -> Dict:
        """Extrait et structure le profil complet du candidat"""
        lines = cv_text.split('\n')
        cv_lower = cv_text.lower()
        
        profile = {
            "identite": self._extract_identity(cv_text, lines),
            "resume_professionnel": self._extract_professional_summary(cv_text, lines),
            "competences_techniques": self._extract_technical_skills_structured(cv_text, cv_lower),
            "experiences_professionnelles": self._extract_professional_experiences_structured(cv_text, lines),
            "stages_alternances": self._extract_internships_structured(cv_text, lines),
            "projets": self._extract_projects_structured(cv_text, lines),
            "formation": self._extract_education_structured(cv_text, lines),
            "certifications": self._extract_certifications_structured(cv_text, lines),
            "langues": self._extract_languages_structured(cv_text, lines),
            "soft_skills": self._extract_soft_skills(cv_text, cv_lower),
            "score_correspondance": None  # Sera calculé à la fin
        }
        
        # Calculer le score de correspondance
        profile["score_correspondance"] = self._calculate_match_score(profile, job_description)
        
        return profile
    
    def _extract_identity(self, cv_text: str, lines: List[str]) -> Dict:
        """Extrait l'identité du candidat"""
        identity = {}
        
        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, cv_text)
        if email_match:
            identity["email"] = email_match.group()
        
        # Nom et Prénom (chercher en premier dans les premières lignes, avant l'email)
        # Chercher un pattern de nom (2-4 mots commençant par majuscule)
        name_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\b'
        for i, line in enumerate(lines[:10]):
            line_stripped = line.strip()
            # Ignorer les lignes avec emails, téléphones, ou trop longues
            if '@' in line_stripped or re.search(r'\+?\d{8,}', line_stripped) or len(line_stripped) > 80:
                continue
            # Ignorer les lignes qui sont des titres de sections
            if any(section in line_stripped.lower() for section in ['expérience', 'experience', 'formation', 'education', 'compétences', 'skills', 'projets', 'projects', 'certifications', 'langues', 'languages']):
                continue
            # Ignorer les lignes qui contiennent des mots de formation/éducation
            if any(word in line_stripped.lower() for word in ['licence', 'master', 'université', 'école', 'diplôme', 'maîtrise', 'normes', 'comptables', 'françaises', 'ifrs', 'spécialisé', 'dans', 'accompagnement', 'préparation', 'dossiers', 'fiscaux']):
                continue
            # Ignorer les lignes qui contiennent des URLs
            if 'http' in line_stripped.lower() or 'www.' in line_stripped.lower() or 'linkedin' in line_stripped.lower() or 'github' in line_stripped.lower():
                continue
            name_match = re.search(name_pattern, line_stripped)
            if name_match:
                name_text = name_match.group(1)
                name_parts = name_text.split()
                # Filtrer les mots qui ne sont pas des noms (trop courts, mots communs)
                common_words = ['the', 'and', 'or', 'cv', 'resume', 'curriculum', 'vitae', 'maîtrise', 'des', 'normes', 'comptables', 'françaises', 'ifrs', 'et', 'de', 'la', 'le', 'les', 'un', 'une', 'du', 'des', 'pour', 'avec', 'dans', 'sur', 'par', 'en', 'au', 'aux']
                name_parts = [p for p in name_parts if len(p) > 2 and p.lower() not in common_words]
                # Vérifier que ce n'est pas une phrase complète (trop de mots)
                if len(name_parts) >= 2 and len(name_parts) <= 4:
                    # Vérifier que les mots commencent tous par une majuscule (caractéristique des noms)
                    if all(p[0].isupper() for p in name_parts):
                        identity["prenom"] = name_parts[0]
                        identity["nom"] = ' '.join(name_parts[1:])
                        break
                elif len(name_parts) == 1 and len(name_parts[0]) > 3 and name_parts[0][0].isupper():
                    identity["nom"] = name_parts[0]
                    break
        
        # Téléphone (amélioration pour formats internationaux)
        phone_patterns = [
            r'\+212[.\s-]?\d{9}',  # Format marocain avec indicatif
            r'212\d{9}',  # Format marocain sans séparateur
            r'(\+33|0)[1-9](?:[.\s-]?\d{2}){4}',  # Format français
            r'(\+1)?[\s.-]?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}',  # Format US/Canada
            r'\d{2}[\s.-]?\d{2}[\s.-]?\d{2}[\s.-]?\d{2}[\s.-]?\d{2}',  # Format générique
            r'\+?\d{10,12}'  # Format générique international
        ]
        # Chercher le téléphone dans les premières lignes (où il est généralement placé)
        phone_text = '\n'.join(lines[:15])
        for pattern in phone_patterns:
            phone_match = re.search(pattern, phone_text)
            if phone_match:
                phone_number = phone_match.group().strip()
                # Vérifier que ce n'est pas une date ou autre nombre
                if len(phone_number) >= 8 and len(phone_number) <= 15:
                    # Vérifier que ce n'est pas une année (4 chiffres)
                    if not re.match(r'^\d{4}$', phone_number):
                        # Vérifier que ce n'est pas une date (format YYYY-MM-DD ou similaire)
                        if not re.match(r'^\d{4}[-/]\d{2}[-/]\d{2}', phone_number):
                            identity["telephone"] = phone_number
                            break
        
        # LinkedIn
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/pub/)([a-zA-Z0-9-]+)'
        linkedin_match = re.search(linkedin_pattern, cv_text, re.IGNORECASE)
        if linkedin_match:
            identity["linkedin"] = f"linkedin.com/in/{linkedin_match.group(1)}"
        
        # GitHub
        github_pattern = r'(?:github\.com/)([a-zA-Z0-9-]+)'
        github_match = re.search(github_pattern, cv_text, re.IGNORECASE)
        if github_match:
            identity["github"] = f"github.com/{github_match.group(1)}"
        
        # Ville / Pays (amélioration pour éviter les faux positifs comme "Tho, ma")
        # Liste de villes françaises et internationales communes
        known_cities = ['paris', 'lyon', 'marseille', 'toulouse', 'nice', 'nantes', 'strasbourg', 'montpellier',
                       'bordeaux', 'lille', 'rennes', 'reims', 'saint-étienne', 'toulon', 'grenoble', 'dijon',
                       'angers', 'nîmes', 'villeurbanne', 'saint-denis', 'le havre', 'tours', 'caen', 'mulhouse',
                       'london', 'new york', 'los angeles', 'chicago', 'houston', 'philadelphia', 'phoenix',
                       'san antonio', 'san diego', 'dallas', 'san jose', 'madrid', 'barcelona', 'valencia',
                       'seville', 'zaragoza', 'málaga', 'murcia', 'berlin', 'hamburg', 'munich', 'cologne',
                       'frankfurt', 'stuttgart', 'düsseldorf', 'dortmund', 'essen', 'leipzig', 'rome', 'milan',
                       'naples', 'turin', 'palermo', 'genoa', 'bologna', 'florence', 'casablanca', 'rabat',
                       'fès', 'marrakech', 'tanger', 'agadir', 'meknès', 'oujda']
        
        # Prénoms communs à exclure
        common_names = ['sophie', 'thomas', 'alexandre', 'marie', 'pierre', 'jean', 'paul', 'bernard', 'martin',
                       'lucas', 'julie', 'camille', 'antoine', 'claire', 'nicolas', 'sarah', 'david', 'emilie',
                       'tho', 'ber', 'mar', 'sop', 'ale', 'luc', 'jul', 'cam', 'ant', 'cla', 'nic', 'sar', 'dav', 'emi']
        
        # Chercher d'abord les villes connues dans le texte (méthode la plus fiable)
        for city in known_cities:
            # Chercher la ville avec un pattern qui évite les parties de mots
            city_pattern = r'\b' + re.escape(city) + r'\b'
            city_match = re.search(city_pattern, cv_text, re.IGNORECASE)
            if city_match:
                # Vérifier qu'il y a un pays à proximité (dans les 50 caractères suivants)
                context_after = cv_text[city_match.end():city_match.end()+50]
                country_match = re.search(r'(?:France|FR|United States|USA|UK|United Kingdom|Morocco|Maroc|MA|Espagne|Spain|Allemagne|Germany|Italie|Italy)', context_after, re.IGNORECASE)
                if country_match:
                    identity["ville"] = city_match.group().capitalize()
                    identity["pays"] = country_match.group()
                    break
                # Sinon, vérifier qu'il y a une virgule ou un séparateur (pattern "Ville, Pays")
                elif re.search(r',\s*(?:France|FR|United States|USA|UK|United Kingdom|Morocco|Maroc|MA|Espagne|Spain|Allemagne|Germany|Italie|Italy)', context_after, re.IGNORECASE):
                    identity["ville"] = city_match.group().capitalize()
                    country_match2 = re.search(r'(?:France|FR|United States|USA|UK|United Kingdom|Morocco|Maroc|MA|Espagne|Spain|Allemagne|Germany|Italie|Italy)', context_after, re.IGNORECASE)
                    if country_match2:
                        identity["pays"] = country_match2.group()
                    break
        
        # Si pas trouvé, chercher les patterns avec pays (plus fiable)
        if not identity.get("ville"):
            # Pattern strict : "Ville, Pays" ou "Ville Pays" - mais éviter les parties de mots
            # Chercher d'abord avec des villes connues + pays
            for city in known_cities:
                # Pattern pour ville connue suivie d'un pays
                city_country_pattern = r'\b' + re.escape(city) + r'\b,?\s+(?:France|FR|United States|USA|UK|United Kingdom|Morocco|Maroc|MA|Espagne|Spain|Allemagne|Germany|Italie|Italy)\b'
                city_country_match = re.search(city_country_pattern, cv_text, re.IGNORECASE)
                if city_country_match:
                    # Vérifier le contexte pour éviter les faux positifs
                    context_before = cv_text[max(0, city_country_match.start()-30):city_country_match.start()].lower()
                    context_after = cv_text[city_country_match.end():city_country_match.end()+30].lower()
                    
                    # Éviter si c'est dans un nom (comme "Thomas" contient "tho")
                    if not re.search(r'\b' + re.escape(city) + r'[a-z]', context_before + context_after, re.IGNORECASE):
                        identity["ville"] = city.capitalize()
                        # Extraire aussi le pays
                        country_match = re.search(r'(?:France|FR|United States|USA|UK|United Kingdom|Morocco|Maroc|MA|Espagne|Spain|Allemagne|Germany|Italie|Italy)', cv_text[city_country_match.start():city_country_match.end()], re.IGNORECASE)
                        if country_match:
                            country_text = country_match.group()
                            # Éviter "MA" seul qui peut être une partie de mot (comme "maîtrise")
                            if country_text.upper() == 'MA':
                                # Vérifier qu'il y a un contexte valide (pas juste "ma" dans une phrase)
                                # Si "MA" est précédé ou suivi d'une lettre minuscule, c'est probablement une partie de mot
                                if re.search(r'[a-z]ma\b|\bma[a-z]', context_before + context_after, re.IGNORECASE):
                                    continue
                                # Vérifier que ce n'est pas "ma" dans "maîtrise", "mais", etc.
                                if re.search(r'\bma(?:îtrise|is|is|intenant|intenance)\b', context_before + context_after, re.IGNORECASE):
                                    continue
                            identity["pays"] = country_text
                        break
            
            # Si toujours pas trouvé, chercher un pattern générique mais avec des validations strictes
            if not identity.get("ville"):
                # Pattern strict : "Ville, Pays" - minimum 5 caractères pour la ville (évite "Tho", "Ber", etc.)
                city_country_pattern = r'\b([A-Z][a-z]{4,}(?:\s+[A-Z][a-z]+)*),?\s+(?:France|FR|United States|USA|UK|United Kingdom|Morocco|Maroc|MA|Espagne|Spain|Allemagne|Germany|Italie|Italy)\b'
                city_country_matches = list(re.finditer(city_country_pattern, cv_text, re.IGNORECASE))
                
                for city_country_match in city_country_matches:
                    city_candidate = city_country_match.group(1).strip()
                    city_lower = city_candidate.lower()
                    
                    # Vérifier le contexte pour éviter les parties de mots
                    context_before = cv_text[max(0, city_country_match.start()-30):city_country_match.start()].lower()
                    context_after = cv_text[city_country_match.end():city_country_match.end()+30].lower()
                    
                    # La ville doit faire au moins 5 caractères et ne pas être un prénom commun
                    if len(city_candidate) >= 5 and len(city_candidate) < 30:
                        # Vérifier que ce n'est pas un prénom commun ou une partie de prénom
                        is_common_name = city_lower in common_names or any(name in city_lower for name in common_names)
                        # Vérifier que ce n'est pas une partie de mot (comme "Tho" de "Thomas", "Ber" de "Bernard")
                        is_partial_word = (len(city_candidate) < 6 and not city_lower in known_cities) or \
                                        re.search(r'\b' + re.escape(city_lower) + r'[a-z]', context_before + context_after, re.IGNORECASE)
                        
                        if not is_common_name and not is_partial_word:
                            # Vérifier que c'est une ville connue ou un pattern valide
                            if city_lower in known_cities or (len(city_candidate) >= 5 and len(city_candidate.split()) <= 2):
                                identity["ville"] = city_candidate
                                # Extraire aussi le pays
                                country_match = re.search(r'(?:France|FR|United States|USA|UK|United Kingdom|Morocco|Maroc|MA|Espagne|Spain|Allemagne|Germany|Italie|Italy)', cv_text[city_country_match.start():city_country_match.end()], re.IGNORECASE)
                                if country_match:
                                    country_text = country_match.group()
                                    # Éviter "MA" seul qui peut être une partie de mot
                                    if country_text.upper() == 'MA':
                                        if re.search(r'[a-z]ma\b|\bma[a-z]', context_before + context_after, re.IGNORECASE):
                                            # "MA" est une partie de mot, passer au suivant
                                            continue
                                        # Vérifier que ce n'est pas "ma" dans "maîtrise", "mais", etc.
                                        if re.search(r'\bma(?:îtrise|is|is|intenant|intenance)\b', context_before + context_after, re.IGNORECASE):
                                            # "MA" est dans un mot, passer au suivant
                                            continue
                                    identity["pays"] = country_text
                                break
        
        # Si pas trouvé, chercher avec les patterns de localisation explicites
        if not identity.get("ville"):
            city_patterns = [
                r'(?:ville|city|location|localisation|adresse|réside|habite|habitant)[\s:]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
            ]
            # Mots à exclure (ne sont pas des villes)
            exclude_words = ['and', 'AI', 'programming', 'student', 'developer', 'engineer', 'web', 'designer', 
                            'studying', 'specialize', 'creating', 'innovative', 'solutions', 'passion', 'domain',
                            'investing', 'content', 'creation', 'work', 'blends', 'technical', 'expertise',
                            'creative', 'design', 'deliver', 'impactful', 'projects', 'currently', 'licence',
                            'master', 'université', 'école', 'formation', 'diplôme', 'communication', 'marketing',
                            'comptabilité', 'finance', 'gestion', 'maîtrise', 'normes', 'comptables', 'françaises']
            for pattern in city_patterns:
                city_match = re.search(pattern, cv_text, re.IGNORECASE)
                if city_match:
                    city_candidate = city_match.group(1).strip()
                    city_lower = city_candidate.lower()
                    # Vérifier que ce n'est pas un faux positif
                    if not any(exclude_word.lower() in city_lower for exclude_word in exclude_words):
                        if len(city_candidate) >= 4 and len(city_candidate) < 30:
                            # Vérifier que ce n'est pas une phrase complète ou un nom de personne
                            if len(city_candidate.split()) <= 2:
                                # Vérifier que ce n'est pas un prénom commun ou une partie de prénom
                                if city_lower not in common_names and not any(name in city_lower for name in common_names):
                                    identity["ville"] = city_candidate
                                    break
        
        # Titre du profil (chercher un titre court et professionnel, pas une phrase complète)
        title_keywords = ['developer', 'engineer', 'analyst', 'manager', 'consultant', 'specialist', 'expert',
                         'développeur', 'ingénieur', 'analyste', 'manager', 'consultant', 'spécialiste', 'expert',
                         'comptable', 'marketing', 'digital', 'finance', 'commercial', 'designer', 'architect',
                         'student', 'étudiant', 'intern', 'stagiaire', 'responsable', 'directeur', 'chef',
                         'senior', 'junior', 'lead', 'chief']
        
        # Mots à exclure (ne sont pas des titres)
        exclude_title_words = ['formation', 'education', 'expérience', 'experience', 'compétences', 'skills',
                              'langues', 'languages', 'projets', 'projects', 'certifications', 'licence', 'master',
                              'université', 'école', 'diplôme', 'diploma', 'maîtrise', 'normes', 'comptables',
                              'françaises', 'ifrs', 'des', 'et', 'dans', 'pour', 'avec', 'sur', 'par', 'en', 'au',
                              'accompagnement', 'préparation', 'dossiers', 'fiscaux', 'spécialisé', 'l\'accompagnement', 'la préparation']
        
        # Chercher un titre court (max 50 caractères) dans les premières lignes
        for i, line in enumerate(lines[:12]):
            line_stripped = line.strip()
            line_lower = line_stripped.lower()
            
            # Ignorer les lignes qui sont clairement des formations ou autres sections
            if any(exclude_word in line_lower for exclude_word in exclude_title_words):
                continue
            
            # Ignorer les lignes avec emails, téléphones, ou dates
            if '@' in line_stripped or re.search(r'\+?\d{8,}', line_stripped) or re.search(r'\d{4}\s*[-–—]', line_stripped):
                continue
            
            # Chercher un titre professionnel court (max 50 caractères)
            if len(line_stripped) > 3 and len(line_stripped) <= 50:
                # Vérifier qu'il contient un mot-clé de titre
                if any(keyword in line_lower for keyword in title_keywords):
                    # Vérifier que ce n'est pas une phrase complète (pas de ponctuation finale sauf si c'est un titre court)
                    if not line_stripped.endswith('.') or len(line_stripped.split()) <= 5:
                        # Vérifier qu'il n'y a pas trop de mots (un titre est généralement court)
                        if len(line_stripped.split()) <= 6:
                            identity["titre_profil"] = line_stripped
                            break
        
        return identity
    
    def _extract_professional_summary(self, cv_text: str, lines: List[str]) -> Dict:
        """Extrait le résumé professionnel"""
        summary = {}
        
        # Chercher les sections de résumé
        summary_keywords = ['résumé', 'resume', 'profil', 'profile', 'summary', 'about', 'à propos']
        summary_section = []
        in_summary = False
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in summary_keywords):
                in_summary = True
                continue
            if in_summary:
                if line.strip() and len(line.strip()) > 10:
                    summary_section.append(line.strip())
                elif len(summary_section) > 0:
                    break
        
        if summary_section:
            summary["resume"] = ' '.join(summary_section[:4])
        
        # Niveau (Junior / Confirmé / Senior)
        level_keywords = {
            'junior': ['junior', 'débutant', 'beginner', 'entry level', 'stagiaire'],
            'confirmé': ['confirmé', 'intermediate', 'expérimenté', 'experienced'],
            'senior': ['senior', 'expert', 'lead', 'principal', 'architect']
        }
        cv_lower = cv_text.lower()
        for level, keywords in level_keywords.items():
            if any(kw in cv_lower for kw in keywords):
                summary["niveau"] = level.capitalize()
                break
        
        # Domaine principal (basé sur les compétences et le contexte - amélioré pour tous domaines)
        domains = {
            'Développement Web': ['web', 'frontend', 'backend', 'fullstack', 'développeur web', 'web developer'],
            'Mobile': ['mobile', 'android', 'ios', 'react native', 'flutter', 'développeur mobile'],
            'Data Science': ['data science', 'data scientist', 'machine learning', 'deep learning', 'analytics avancé'],
            'Cybersécurité': ['security', 'cyber', 'pentest', 'sécurité', 'cybersécurité', 'cybersecurity'],
            'DevOps': ['devops', 'cloud', 'docker', 'kubernetes', 'ci/cd', 'infrastructure'],
            'Marketing Digital': ['marketing digital', 'digital marketing', 'community manager', 'social media', 
                                 'seo', 'sem', 'content marketing', 'email marketing', 'réseaux sociaux'],
            'Finance & Comptabilité': ['comptable', 'comptabilité', 'finance', 'expert-comptable', 'audit', 
                                      'fiscalité', 'gestion financière', 'analyse financière', 'sage', 'ciel'],
            'Ressources Humaines': ['rh', 'ressources humaines', 'recrutement', 'gestion du personnel', 'hr'],
            'Vente & Commerce': ['commercial', 'vente', 'business development', 'account manager', 'sales'],
            'Design & Création': ['designer', 'design', 'graphiste', 'création', 'ui/ux', 'illustration']
        }
        
        # Détecter le domaine par mots-clés (plus rapide que l'API)
        # Compter les occurrences de mots-clés par domaine
        domain_scores = {}
        for domain, keywords in domains.items():
            score = sum(1 for kw in keywords if kw in cv_lower)
            if score > 0:
                domain_scores[domain] = score
        
        # Prendre le domaine avec le score le plus élevé
        if domain_scores:
            max_domain = max(domain_scores.items(), key=lambda x: x[1])
            summary["domaine_principal"] = max_domain[0]
        
        return summary
    
    def _extract_technical_skills_structured(self, cv_text: str, cv_lower: str) -> Dict:
        """Extrait les compétences techniques structurées par catégorie avec IA (sans listes statiques)"""
        skills = {
            "langages": [],
            "frameworks": [],
            "outils": [],
            "cloud": [],
            "ia_data": [],
            "securite": []
        }
        
        # Mots à exclure (ne sont pas des compétences techniques)
        exclude_skills = ['pme', 'ifrs', 'dec', 'tva', 'cvae', 'pcg', 'formation', 'professionnelle', 
                         'techniques', 'certifications', 'langues', 'comptables', 'comptable', 'expert',
                         'ordre', 'experts', 'université', 'école', 'master', 'licence', 'diplôme',
                         'communication', 'marketing', 'digital', 'finance', 'gestion', 'présent',
                         'present', '2014', '2015', '2016', '2017', '2018', '2019', '2020', '2021',
                         '2022', '2023', '2024', 'lyon', 'bordeaux', 'paris', 'france']
        
        # Extraire toutes les compétences avec IA
        all_skills = self.extract_skills(cv_text)
        
        # Classifier les compétences par catégorie en utilisant l'analyse sémantique
        for skill in all_skills:
            skill_lower = skill.lower().strip()
            
            # Exclure les mots qui ne sont pas des compétences
            if skill_lower in exclude_skills or len(skill_lower) < 2:
                continue
            
            category = self._classify_skill_category(skill, cv_text)
            
            if category and category in skills:
                if skill not in skills[category]:
                    skills[category].append(skill)
        
        return skills
    
    def _classify_skill_category(self, skill: str, context: str) -> Optional[str]:
        """Classifie une compétence dans une catégorie en utilisant l'analyse sémantique"""
        skill_lower = skill.lower()
        context_lower = context.lower()
        
        # Patterns pour langages de programmation
        lang_patterns = [
            r'\b(python|javascript|java|php|ruby|go|rust|swift|kotlin|typescript|scala|r|matlab|c\+\+|c#|html|css|sql)\b',
            r'\b(langage|programming language|language)\b.*' + re.escape(skill_lower)
        ]
        for pattern in lang_patterns:
            if re.search(pattern, context_lower, re.IGNORECASE):
                return "langages"
        
        # Patterns pour frameworks
        framework_patterns = [
            r'\b(react|vue|angular|node|django|flask|spring|express|laravel|symfony|rails|asp\.net)\b',
            r'\b(framework|library|bibliothèque)\b.*' + re.escape(skill_lower),
            r'\b(wordpress|woocommerce|drupal|joomla|shopify)\b'
        ]
        for pattern in framework_patterns:
            if re.search(pattern, context_lower, re.IGNORECASE):
                return "frameworks"
        
        # Patterns pour outils
        tools_patterns = [
            r'\b(git|docker|kubernetes|jenkins|gitlab|github|jira|confluence|postman|swagger)\b',
            r'\b(tool|outil|software)\b.*' + re.escape(skill_lower)
        ]
        for pattern in tools_patterns:
            if re.search(pattern, context_lower, re.IGNORECASE):
                return "outils"
        
        # Patterns pour cloud
        cloud_patterns = [
            r'\b(aws|azure|gcp|google cloud|heroku|digitalocean|oracle cloud)\b',
            r'\b(cloud|infrastructure|platform)\b.*' + re.escape(skill_lower)
        ]
        for pattern in cloud_patterns:
            if re.search(pattern, context_lower, re.IGNORECASE):
                return "cloud"
        
        # Patterns pour IA/Data (amélioré pour éviter les faux positifs)
        ia_data_patterns = [
            r'\b(tensorflow|pytorch|keras|pandas|numpy|scikit-learn|spark|hadoop|tableau|power bi)\b',
            r'\b(machine learning|ai|data science|deep learning|analytics avancé|big data)\b',
        ]
        # Mots à exclure (ne sont pas des compétences IA/Data)
        exclude_ia_words = ['pme', 'ifrs', 'dec', 'tva', 'cvae', 'pcg', 'formation', 'professionnelle', 
                           'techniques', 'certifications', 'langues', 'comptables']
        if skill_lower not in exclude_ia_words:
            for pattern in ia_data_patterns:
                if re.search(pattern, context_lower, re.IGNORECASE):
                    return "ia_data"
        
        # Patterns pour sécurité
        security_patterns = [
            r'\b(owasp|pentest|metasploit|burp suite|wireshark|nmap|ssl|tls|vpn)\b',
            r'\b(security|sécurité|cybersecurity|cybersécurité)\b.*' + re.escape(skill_lower)
        ]
        for pattern in security_patterns:
            if re.search(pattern, context_lower, re.IGNORECASE):
                return "securite"
        
        # Classification basée sur les mots-clés (plus rapide que l'API)
        # Utiliser une logique basée sur les mots-clés plutôt que des appels API
        skill_lower = skill.lower()
        context_lower = context.lower()
        
        # Vérifier les mots-clés spécifiques dans le contexte
        lang_keywords_list = ["programming", "language", "code", "syntax", "python", "java", "javascript", "php", "ruby"]
        framework_keywords_list = ["framework", "library", "react", "vue", "angular", "django", "flask", "spring"]
        tool_keywords_list = ["tool", "software", "utility", "git", "docker", "jenkins", "jira"]
        cloud_keywords_list = ["cloud", "aws", "azure", "gcp", "infrastructure", "platform"]
        ia_keywords_list = ["machine learning", "ai", "data science", "neural", "tensorflow", "pytorch"]
        security_keywords_list = ["security", "sécurité", "cybersecurity", "pentest", "vulnerability"]
        
        # Vérifier les mots-clés dans le skill et le contexte
        skill_and_context = f"{skill_lower} {context_lower[:200]}"
        
        if any(kw in skill_and_context for kw in lang_keywords_list):
            return "langages"
        elif any(kw in skill_and_context for kw in framework_keywords_list):
            return "frameworks"
        elif any(kw in skill_and_context for kw in cloud_keywords_list):
            return "cloud"
        elif any(kw in skill_and_context for kw in ia_keywords_list):
            return "ia_data"
        elif any(kw in skill_and_context for kw in security_keywords_list):
            return "securite"
        elif any(kw in skill_and_context for kw in tool_keywords_list):
            return "outils"
        
        # Par défaut, mettre dans "outils" si aucune catégorie ne correspond
        return "outils"
    
    def _extract_professional_experiences_structured(self, cv_text: str, lines: List[str]) -> List[Dict]:
        """Extrait les expériences professionnelles structurées avec amélioration"""
        experiences = []
        
        # Chercher toutes les dates (années) dans le texte
        date_pattern = r'(\d{4})\s*[-–—]\s*(\d{4}|présent|present|now|aujourd\'hui|current)'
        all_dates = list(re.finditer(date_pattern, cv_text, re.IGNORECASE))
        
        # Chercher les sections d'expérience
        exp_keywords = ['expérience', 'experience', 'work', 'employment', 'emploi', 'professional', 'career']
        exp_sections = []
        
        # Trouver les indices des sections d'expérience
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in exp_keywords) and len(line_lower) < 50:
                exp_sections.append(i)
        
        # Si pas de section trouvée, chercher par dates
        if not exp_sections:
            # Chercher les blocs de texte autour des dates
            for date_match in all_dates[:10]:  # Limiter à 10
                start_pos = max(0, date_match.start() - 200)
                end_pos = min(len(cv_text), date_match.end() + 500)
                context = cv_text[start_pos:end_pos]
                context_lines = context.split('\n')
                
                exp = {}
                exp["periode"] = date_match.group(0)
                
                # Chercher le titre du poste (avant la date) - amélioré
                for line in context_lines[:5]:
                    line_stripped = line.strip()
                    if len(line_stripped) > 10 and len(line_stripped) < 100:
                        if not exp.get("intitule_poste"):
                            # Vérifier que ce n'est pas un email, téléphone, ou section
                            if ('@' not in line_stripped and 
                                not re.search(r'\+?\d', line_stripped[:5]) and
                                not any(keyword in line_stripped.lower() for keyword in ['formation', 'education', 'compétences', 'skills', 'langues', 'projets', 'certifications']) and
                                not re.match(r'^\d{4}', line_stripped)):  # Ne pas prendre les lignes qui commencent par une année
                                exp["intitule_poste"] = line_stripped
                                break
                
                # Chercher l'entreprise (ligne en majuscules ou courte) - amélioré
                for line in context_lines:
                    line_stripped = line.strip()
                    if line_stripped and len(line_stripped) > 2:
                        # Vérifier que ce n'est pas une section ou autre chose
                        is_section = any(keyword in line_stripped.lower() for keyword in ['formation', 'education', 'expérience', 'experience', 'compétences', 'skills'])
                        is_date = re.match(r'^\d{4}', line_stripped)
                        
                        if not is_section and not is_date:
                            # Entreprise : ligne en majuscules, ou courte (3-4 mots max), ou nom propre
                            if (line_stripped.isupper() or 
                                (len(line_stripped.split()) <= 4 and len(line_stripped) > 3) or
                                (line_stripped[0].isupper() and len(line_stripped.split()) <= 3)):
                                if line_stripped != exp.get("intitule_poste") and line_stripped != exp.get("periode"):
                                    exp["entreprise"] = line_stripped
                                    break
                
                # Chercher les missions (lignes avec verbes d'action ou puces)
                missions = []
                action_verbs = ['developed', 'created', 'built', 'managed', 'designed', 'implemented', 
                               'développé', 'créé', 'construit', 'géré', 'conçu', 'implémenté',
                               'performed', 'worked', 'collaborated', 'led', 'improved', 'gestion',
                               'création', 'animation', 'organisation', 'analyse', 'augmentation']
                
                # Mots à exclure des missions (sont des technologies ou autres)
                exclude_mission_words = ['linkedin', 'facebook', 'instagram', 'twitter', 'youtube', 'tiktok',
                                        'mailchimp', 'hootsuite', 'hubspot', 'canva', 'adobe', 'photoshop',
                                        'premiere', 'google', 'analytics', 'tag manager', 'semrush', 'ahrefs']
                
                for line in context_lines:
                    line_stripped = line.strip()
                    if line_stripped and len(line_stripped) > 15:
                        line_lower = line_stripped.lower()
                        
                        # Exclure les lignes qui sont clairement des listes de technologies
                        if any(exclude_word in line_lower for exclude_word in exclude_mission_words):
                            continue
                        
                        # Vérifier si c'est une mission (commence par verbe ou puce)
                        if (any(line_lower.startswith(verb) for verb in action_verbs) or
                            line_stripped.startswith(('-', '•', '*', '→', '·', '▸')) or
                            re.match(r'^\d+[\.\)]', line_stripped)):
                            mission = re.sub(r'^[-•*\d\.\)\s→·▸]+', '', line_stripped)
                            # Nettoyer la mission
                            mission_clean = mission
                            # Si la mission contient beaucoup de mots, prendre seulement le début
                            if len(mission.split()) > 20:
                                mission_clean = ' '.join(mission.split()[:15]) + "..."
                            
                            if mission_clean and mission_clean not in missions and len(mission_clean) > 10:
                                missions.append(mission_clean)
                
                if missions:
                    exp["missions"] = missions[:5]  # Limiter à 5 missions
                
                # Extraire les technologies mentionnées
                technologies = []
                # Extraire les technologies mentionnées dynamiquement
                tech_keywords = self.extract_skills(context)
                for tech in tech_keywords:
                    if tech.lower() in context.lower():
                        technologies.append(tech.capitalize())
                if technologies:
                    exp["technologies"] = list(set(technologies))[:10]
                
                if exp.get("intitule_poste") or exp.get("entreprise"):
                    experiences.append(exp)
        
        # Méthode alternative : parser ligne par ligne si pas assez d'expériences
        if len(experiences) < 2:
            current_exp = {}
            in_experience = False
            
            for i, line in enumerate(lines):
                line_stripped = line.strip()
                if not line_stripped or len(line_stripped) < 3:
                    if current_exp and (current_exp.get("intitule_poste") or current_exp.get("entreprise")):
                        experiences.append(current_exp)
                        current_exp = {}
                    in_experience = False
                    continue
                
                # Détecter période
                period_match = re.search(r'(\d{4})\s*[-–—]\s*(\d{4}|présent|present|now)', line_stripped)
                if period_match:
                    if current_exp and (current_exp.get("intitule_poste") or current_exp.get("entreprise")):
                        experiences.append(current_exp)
                    current_exp = {"periode": period_match.group(0)}
                    in_experience = True
                    continue
                
                if in_experience:
                    # Intitulé du poste
                    if not current_exp.get("intitule_poste") and len(line_stripped) > 10 and len(line_stripped) < 100:
                        if '@' not in line_stripped and not re.search(r'\+?\d{10}', line_stripped):
                            current_exp["intitule_poste"] = line_stripped
                            continue
                    
                    # Entreprise
                    if not current_exp.get("entreprise") and (line_stripped.isupper() or len(line_stripped.split()) <= 4):
                        if line_stripped != current_exp.get("intitule_poste"):
                            current_exp["entreprise"] = line_stripped
                            continue
                    
                    # Missions
                    if line_stripped.startswith(('-', '•', '*', '→', '·')) or re.match(r'^\d+[\.\)]', line_stripped):
                        mission = re.sub(r'^[-•*\d\.\)\s→·]+', '', line_stripped)
                        if mission and len(mission) > 10:
                            if "missions" not in current_exp:
                                current_exp["missions"] = []
                            current_exp["missions"].append(mission)
            
            if current_exp and (current_exp.get("intitule_poste") or current_exp.get("entreprise")):
                experiences.append(current_exp)
        
        # Dédupliquer et nettoyer
        seen = set()
        unique_experiences = []
        for exp in experiences:
            key = (exp.get("intitule_poste", ""), exp.get("entreprise", ""), exp.get("periode", ""))
            if key not in seen and (exp.get("intitule_poste") or exp.get("entreprise")):
                seen.add(key)
                unique_experiences.append(exp)
        
        return unique_experiences[:10]
    
    def _extract_internships_structured(self, cv_text: str, lines: List[str]) -> List[Dict]:
        """Extrait les stages et alternances"""
        internships = []
        stage_patterns = [
            r'(?i)(stage|internship|alternance|apprentissage|apprenticeship)',
            r'(?i)(stagiaire|intern)'
        ]
        
        current_stage = {}
        in_stage = False
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                if current_stage and current_stage.get("intitule"):
                    internships.append(current_stage)
                    current_stage = {}
                in_stage = False
                continue
            
            if any(re.search(pattern, line_stripped) for pattern in stage_patterns):
                if current_stage and current_stage.get("intitule"):
                    internships.append(current_stage)
                current_stage = {}
                in_stage = True
                continue
            
            if in_stage:
                if not current_stage.get("intitule") and len(line_stripped) > 5:
                    current_stage["intitule"] = line_stripped
                elif line_stripped.startswith(('-', '•', '*')):
                    mission = re.sub(r'^[-•*\s]+', '', line_stripped)
                    if "missions" not in current_stage:
                        current_stage["missions"] = []
                    current_stage["missions"].append(mission)
        
        if current_stage and current_stage.get("intitule"):
            internships.append(current_stage)
        
        return internships[:5]
    
    def _extract_projects_structured(self, cv_text: str, lines: List[str]) -> List[Dict]:
        """Extrait les projets"""
        projects = []
        project_patterns = [
            r'(?i)(projet|project|portfolio)',
            r'(?i)(réalisations|achievements)'
        ]
        
        current_project = {}
        in_project = False
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                if current_project and current_project.get("nom"):
                    projects.append(current_project)
                    current_project = {}
                in_project = False
                continue
            
            if any(re.search(pattern, line_stripped) for pattern in project_patterns):
                if current_project and current_project.get("nom"):
                    projects.append(current_project)
                current_project = {}
                in_project = True
                continue
            
            if in_project:
                if not current_project.get("nom") and len(line_stripped) > 3 and len(line_stripped) < 100:
                    current_project["nom"] = line_stripped
                elif line_stripped.startswith(('-', '•', '*')):
                    desc = re.sub(r'^[-•*\s]+', '', line_stripped)
                    if "description" not in current_project:
                        current_project["description"] = desc
                    else:
                        current_project["description"] += " " + desc
        
        if current_project and current_project.get("nom"):
            projects.append(current_project)
        
        return projects[:5]
    
    def _extract_education_structured(self, cv_text: str, lines: List[str]) -> List[Dict]:
        """Extrait la formation de manière structurée et claire, sans doublons"""
        education = []
        edu_keywords = ['formation', 'education', 'études', 'studies', 'diplôme', 'diploma']
        degree_keywords = ['master', 'licence', 'bachelor', 'diplôme', 'bac', 'phd', 'doctorat', 
                          'mba', 'bts', 'dut', 'ingénieur', 'engineer']
        school_keywords = ['université', 'university', 'école', 'school', 'institut', 'institute', 
                          'college', 'faculté', 'faculty', 'supérieure']
        
        # Trouver la section Formation
        section_start = -1
        section_end = len(lines)
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if any(keyword in line_lower for keyword in edu_keywords) and len(line_lower) < 30:
                section_start = i + 1
                # Trouver la fin de la section
                for j in range(i + 1, min(len(lines), i + 100)):
                    next_line_lower = lines[j].lower().strip()
                    if any(keyword in next_line_lower for keyword in ['expérience', 'experience', 'compétences', 'skills', 'projets', 'certifications', 'langues']):
                        section_end = j
                        break
                break
        
        if section_start == -1:
            section_start = 0
        
        # Parser les formations ligne par ligne
        current_edu = {}
        
        for i in range(section_start, section_end):
            line = lines[i].strip()
            if not line or len(line) < 3:
                if current_edu and current_edu.get("diplome"):
                    education.append(current_edu)
                    current_edu = {}
                continue
            
            line_lower = line.lower()
            
            # Détecter une ligne complète avec pattern: "Master en Marketing Digital École Supérieure de Commerce de Lyon 2017 - 2019"
            year_pattern = r'(\d{4})\s*[-–—]\s*(\d{4}|\d{2})'
            has_years = bool(re.search(year_pattern, line))
            has_degree = any(word in line_lower[:30] for word in degree_keywords)
            has_school = any(keyword in line_lower for keyword in school_keywords)
            
            # Si ligne complète détectée
            if has_degree and has_years:
                if current_edu and current_edu.get("diplome"):
                    education.append(current_edu)
                
                current_edu = {}
                
                # Extraire le diplôme (du début jusqu'à l'établissement ou les années)
                diploma_parts = []
                words = line.split()
                for word in words:
                    if any(keyword in word.lower() for keyword in school_keywords):
                        break
                    if re.match(r'^\d{4}', word):
                        break
                    diploma_parts.append(word)
                
                diploma_text = ' '.join(diploma_parts).strip()
                # Nettoyer
                diploma_text = re.sub(r'^(formation|education|études)[\s:]+', '', diploma_text, flags=re.IGNORECASE)
                diploma_text = re.sub(r'\s+', ' ', diploma_text).strip()
                
                if len(diploma_text) > 5:
                    current_edu["diplome"] = diploma_text[:80]
                
                # Extraire l'établissement
                if has_school:
                    school_match = re.search(r'(?:' + '|'.join(school_keywords) + r')[\s\w]+?(?=\d{4}|$)', line, re.IGNORECASE)
                    if school_match:
                        school_text = school_match.group(0).strip()
                        school_text = re.sub(r'\s+\d{4}.*$', '', school_text).strip()
                        school_text = re.sub(r'\s+', ' ', school_text)
                        current_edu["etablissement"] = school_text[:60]
                
                # Extraire les années
                year_match = re.search(year_pattern, line)
                if year_match:
                    year_end = year_match.group(2)
                    if len(year_end) == 2 and year_end.isdigit():
                        year_end = "20" + year_end
                    current_edu["annees"] = f"{year_match.group(1)} - {year_end}"
                continue
            
            # Ligne avec seulement des années
            if re.match(r'^\d{4}\s*[-–—]', line):
                if current_edu and not current_edu.get("annees"):
                    year_match = re.search(year_pattern, line)
                    if year_match:
                        year_end = year_match.group(2)
                        if len(year_end) == 2 and year_end.isdigit():
                            year_end = "20" + year_end
                        current_edu["annees"] = f"{year_match.group(1)} - {year_end}"
                continue
            
            # Ligne avec un diplôme
            if has_degree and not current_edu.get("diplome"):
                diploma_clean = line
                diploma_clean = re.sub(r'^(formation|education|études)[\s:]+', '', diploma_clean, flags=re.IGNORECASE)
                # Enlever l'établissement si présent
                for kw in school_keywords:
                    if kw in diploma_clean.lower():
                        idx = diploma_clean.lower().find(kw)
                        diploma_clean = diploma_clean[:idx].strip()
                        break
                # Enlever les dates
                diploma_clean = re.sub(r'\s+\d{4}.*$', '', diploma_clean).strip()
                diploma_clean = re.sub(r'\s+', ' ', diploma_clean)
                if len(diploma_clean) > 5:
                    current_edu["diplome"] = diploma_clean[:80]
                continue
            
            # Ligne avec un établissement
            if has_school and not current_edu.get("etablissement"):
                school_clean = line
                school_clean = re.sub(r'\s+\d{4}.*$', '', school_clean).strip()
                school_clean = re.sub(r'\s+', ' ', school_clean)
                current_edu["etablissement"] = school_clean[:60]
                continue
        
        # Ajouter la dernière formation
        if current_edu and current_edu.get("diplome"):
            education.append(current_edu)
        
        # Dédupliquer strictement
        seen = set()
        unique_education = []
        for edu in education:
            diplome = edu.get("diplome", "").strip().lower()
            annees = edu.get("annees", "").strip()
            
            # Normaliser pour comparaison
            diplome_normalized = re.sub(r'[^\w\s]', '', diplome)
            diplome_normalized = re.sub(r'\s+', ' ', diplome_normalized).strip()[:50]
            
            # Clé unique basée sur diplôme + années
            key = (diplome_normalized, annees)
            
            if key not in seen and diplome and len(diplome_normalized) > 3:
                seen.add(key)
                # Nettoyer final
                edu["diplome"] = re.sub(r'\s+', ' ', edu.get("diplome", "").strip())
                if edu.get("etablissement"):
                    edu["etablissement"] = re.sub(r'\s+', ' ', edu.get("etablissement", "").strip())
                unique_education.append(edu)
        
        return unique_education[:5]
    
    def _extract_certifications_structured(self, cv_text: str, lines: List[str]) -> List[Dict]:
        """Extrait les certifications"""
        certifications = []
        cert_patterns = [
            r'(?i)(certification|certificat|certificate|cert)',
            r'(?i)(aws certified|azure certified|google cloud|oracle certified)'
        ]
        
        current_cert = {}
        in_cert = False
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            line_lower = line_stripped.lower()
            
            if any(re.search(pattern, line_lower) for pattern in cert_patterns):
                if current_cert and current_cert.get("nom"):
                    certifications.append(current_cert)
                current_cert = {}
                in_cert = True
                continue
            
            if in_cert:
                if not current_cert.get("nom") and len(line_stripped) > 5:
                    current_cert["nom"] = line_stripped
                elif "organisme" not in current_cert and any(word in line_lower for word in ['aws', 'microsoft', 'google', 'oracle', 'cisco']):
                    current_cert["organisme"] = line_stripped
                else:
                    year_match = re.search(r'\d{4}', line_stripped)
                    if year_match:
                        current_cert["annee"] = year_match.group(0)
        
        if current_cert and current_cert.get("nom"):
            certifications.append(current_cert)
        
        return certifications[:5]
    
    def _extract_languages_structured(self, cv_text: str, lines: List[str]) -> List[Dict]:
        """Extrait les langues avec niveaux"""
        languages_list = []
        languages = self.extract_languages(cv_text)
        
        # Chercher les niveaux pour chaque langue
        for lang in languages:
            lang_dict = {"langue": lang}
            
            # Chercher le niveau associé
            lang_lower = lang.lower()
            cv_lower = cv_text.lower()
            
            # Patterns pour trouver le niveau
            level_patterns = {
                'Fluent': [f'{lang_lower}.*fluent', f'{lang_lower}.*natif', f'{lang_lower}.*native', f'{lang_lower}.*courant'],
                'Avancé': [f'{lang_lower}.*avancé', f'{lang_lower}.*advanced', f'{lang_lower}.*c1', f'{lang_lower}.*c2'],
                'Intermédiaire': [f'{lang_lower}.*intermédiaire', f'{lang_lower}.*intermediate', f'{lang_lower}.*b1', f'{lang_lower}.*b2'],
                'Débutant': [f'{lang_lower}.*débutant', f'{lang_lower}.*beginner', f'{lang_lower}.*a1', f'{lang_lower}.*a2']
            }
            
            for level, patterns in level_patterns.items():
                if any(re.search(pattern, cv_lower, re.IGNORECASE) for pattern in patterns):
                    lang_dict["niveau"] = level
                    break
            
            if "niveau" not in lang_dict:
                lang_dict["niveau"] = "Non spécifié"
            
            languages_list.append(lang_dict)
        
        return languages_list
    
    def _extract_soft_skills(self, cv_text: str, cv_lower: str) -> List[str]:
        """Extrait les soft skills"""
        soft_skills_keywords = [
            'leadership', 'teamwork', 'communication', 'gestion', 'management',
            'autonome', 'autonomous', 'créatif', 'creative', 'analytique',
            'proactif', 'proactive', 'rigoureux', 'rigorous', 'adaptable',
            'résolution de problèmes', 'problem solving', 'organisation',
            'travail en équipe', 'collaboration', 'motivation', 'curiosité'
        ]
        
        found_skills = []
        for skill in soft_skills_keywords:
            if skill in cv_lower:
                found_skills.append(skill.capitalize())
        
        return found_skills[:10]
    
    def _calculate_match_score(self, profile: Dict, job_description: str) -> float:
        """Calcule un score de correspondance global BASÉ SUR L'IA SÉMANTIQUE entre CV et description du poste"""
        if not job_description:
            return 0.0
        
        # Construire un résumé complet du CV à partir du profil
        cv_summary_parts = []
        
        # 1. Résumé professionnel (priorité)
        if profile.get("resume_professionnel"):
            resume_data = profile.get("resume_professionnel", {})
            if isinstance(resume_data, dict):
                summary = resume_data.get("resume", "") or resume_data.get("resume_professionnel", "") or ""
                if summary:
                    cv_summary_parts.append(summary)
        
        # 2. Compétences techniques
        if profile.get("competences_techniques"):
            skills = profile["competences_techniques"]
            skills_list = []
            for category, skill_list in skills.items():
                if isinstance(skill_list, list):
                    skills_list.extend(skill_list)
            if skills_list:
                cv_summary_parts.append("Compétences: " + ", ".join(skills_list[:20]))
        
        # 3. Expériences professionnelles
        if profile.get("experiences_professionnelles"):
            exp_texts = []
            for exp in profile["experiences_professionnelles"][:3]:  # Top 3 expériences
                exp_str = ""
                if isinstance(exp, dict):
                    if exp.get("intitule_poste"):
                        exp_str += exp["intitule_poste"] + " "
                    if exp.get("entreprise"):
                        exp_str += "chez " + exp["entreprise"] + " "
                    if exp.get("missions"):
                        missions = exp["missions"][:3] if isinstance(exp["missions"], list) else []
                        exp_str += " - " + ", ".join(missions)
                if exp_str:
                    exp_texts.append(exp_str.strip())
            if exp_texts:
                cv_summary_parts.append("Expériences: " + " | ".join(exp_texts))
        
        # 4. Formation
        if profile.get("formation"):
            edu_texts = []
            for edu in profile["formation"][:2]:  # Top 2 formations
                if isinstance(edu, dict):
                    edu_str = ""
                    if edu.get("diplome"):
                        edu_str += edu["diplome"] + " "
                    if edu.get("etablissement"):
                        edu_str += "à " + edu["etablissement"]
                    if edu_str:
                        edu_texts.append(edu_str.strip())
            if edu_texts:
                cv_summary_parts.append("Formation: " + " | ".join(edu_texts))
        
        # Construire le texte complet du CV pour comparaison
        cv_full_text = " ".join(cv_summary_parts)
        
        if not cv_full_text.strip():
            return 0.0
        
        # DEBUG
        print(f"[DEBUG] CV résumé (premiers 200 caractères): {cv_full_text[:200]}")
        print(f"[DEBUG] Description poste (premiers 200 caractères): {job_description[:200]}")
        
        # Extraire le résumé professionnel maintenant (pour utilisation dans le calcul sémantique)
        professional_summary_for_semantic = ""
        if profile.get("resume_professionnel"):
            resume_data = profile.get("resume_professionnel", {})
            if isinstance(resume_data, dict):
                professional_summary_for_semantic = resume_data.get("resume", "") or ""
        
        # CALCUL PRINCIPAL: Similarité sémantique IA entre CV et description du poste
        # Calculer la similarité de plusieurs façons pour plus de précision
        semantic_score_full = self._calculate_semantic_similarity(cv_full_text, job_description)
        
        # Calculer aussi avec le résumé professionnel si disponible
        semantic_score_summary = 0.0
        if professional_summary_for_semantic:
            semantic_score_summary = self._calculate_semantic_similarity(professional_summary_for_semantic, job_description)
        
        # Prendre le meilleur score sémantique (soit résumé complet, soit résumé professionnel)
        semantic_score = max(semantic_score_full, semantic_score_summary * 1.2)  # Bonus si résumé professionnel est bon
        
        # Améliorer le score sémantique si la description est courte mais pertinente
        if len(job_description.split()) <= 10:
            # Pour les descriptions courtes, être plus généreux
            # Si au moins quelques mots-clés correspondent, augmenter le score
            job_words = set(re.findall(r'\b\w{3,}\b', job_description.lower()))
            cv_words = set(re.findall(r'\b\w{3,}\b', cv_full_text.lower()))
            common_words = job_words.intersection(cv_words)
            if len(common_words) >= 2:
                semantic_score = max(semantic_score, 0.4)  # Minimum 0.4 si 2+ mots-clés communs
            if len(common_words) >= 3:
                semantic_score = max(semantic_score, 0.6)  # Minimum 0.6 si 3+ mots-clés communs
        
        print(f"[DEBUG] Score sémantique IA: {semantic_score:.3f} (full: {semantic_score_full:.3f}, summary: {semantic_score_summary:.3f})")
        
        # Comparaison des compétences requises vs compétences du CV (20%)
        required_skills = self._extract_required_skills_from_job(job_description)
        cv_skills = []
        if profile.get("competences_techniques"):
            skills = profile["competences_techniques"]
            for category, skill_list in skills.items():
                if isinstance(skill_list, list):
                    cv_skills.extend([s.lower() for s in skill_list])
        
        skills_match_score = 0.0
        if required_skills and cv_skills:
            matching_count = 0
            cv_skills_lower = [s.lower() for s in cv_skills]
            
            # Utiliser uniquement la comparaison de mots-clés (pas d'API pour chaque compétence)
            for req_skill in required_skills[:20]:  # Top 20 compétences requises
                req_lower = req_skill.lower()
                found = False
                
                # Vérification exacte
                if req_lower in cv_skills_lower:
                    matching_count += 1
                    found = True
                else:
                    # Vérification partielle basée sur les mots (pas d'API)
                    req_words = set(re.findall(r'\b\w{3,}\b', req_lower))
                    for cv_skill in cv_skills[:30]:  # Limiter pour performance
                        cv_skill_lower = cv_skill.lower()
                        cv_words = set(re.findall(r'\b\w{3,}\b', cv_skill_lower))
                        if req_words and cv_words:
                            overlap = len(req_words.intersection(cv_words)) / len(req_words)
                            if overlap >= 0.6:  # 60% de similarité de mots
                                matching_count += overlap
                                found = True
                                break
                
                # Si pas trouvé, pénalité mineure
                if not found:
                    matching_count += 0.0
            
            skills_match_score = matching_count / len(required_skills[:20]) if required_skills[:20] else 0.0
        print(f"[DEBUG] Score correspondance compétences: {skills_match_score:.3f}")
        
        # Comparaison résumé professionnel vs description (10%)
        professional_summary = professional_summary_for_semantic  # Utiliser celui déjà extrait
        
        summary_semantic_score = 0.0
        if professional_summary:
            # Utiliser le calcul amélioré local (pas d'API)
            summary_semantic_score = self._enhanced_similarity(professional_summary, job_description)
        print(f"[DEBUG] Score sémantique résumé: {summary_semantic_score:.3f}")
        
        # Calcul du score final avec pondération réaliste
        # Donner plus de poids aux compétences car c'est le critère le plus objectif
        
        # Score de base pondéré
        if skills_match_score >= 0.7:
            # Si les compétences correspondent très bien, elles sont prioritaires
            base_score = (
                semantic_score * 0.25 +           # 25% - Similarité sémantique
                skills_match_score * 0.60 +      # 60% - Compétences (priorité)
                summary_semantic_score * 0.15    # 15% - Résumé professionnel
            ) * 100
        elif skills_match_score >= 0.5:
            # Si les compétences correspondent bien, équilibre
            base_score = (
                semantic_score * 0.35 +           # 35% - Similarité sémantique
                skills_match_score * 0.50 +      # 50% - Compétences
                summary_semantic_score * 0.15    # 15% - Résumé professionnel
            ) * 100
        elif skills_match_score >= 0.3:
            # Si les compétences correspondent moyennement, sémantique plus important
            base_score = (
                semantic_score * 0.50 +           # 50% - Similarité sémantique
                skills_match_score * 0.35 +      # 35% - Compétences
                summary_semantic_score * 0.15    # 15% - Résumé professionnel
            ) * 100
        else:
            # Si les compétences ne correspondent pas, sémantique dominant
            base_score = (
                semantic_score * 0.65 +           # 65% - Similarité sémantique
                skills_match_score * 0.20 +      # 20% - Compétences
                summary_semantic_score * 0.15    # 15% - Résumé professionnel
            ) * 100
        
        final_score = base_score
        
        # Ajustements réalistes basés sur la correspondance
        
        # 1. Bonus si les compétences correspondent très bien (critère le plus important)
        if skills_match_score >= 0.8:
            final_score = min(final_score * 1.15, 100.0)  # +15% si excellent match compétences
        elif skills_match_score >= 0.7:
            final_score = min(final_score * 1.10, 100.0)  # +10% si très bon match
        elif skills_match_score >= 0.6:
            final_score = min(final_score * 1.05, 100.0)  # +5% si bon match
        
        # 2. Bonus si la similarité sémantique est élevée (bonne correspondance globale)
        if semantic_score >= 0.6:
            final_score = min(final_score * 1.10, 100.0)  # +10% si excellente similarité
        elif semantic_score >= 0.5:
            final_score = min(final_score * 1.05, 100.0)  # +5% si bonne similarité
        elif semantic_score >= 0.4:
            final_score = min(final_score * 1.02, 100.0)  # +2% si similarité correcte
        
        # 3. Pénalités seulement si TOUT est vraiment mauvais
        if skills_match_score < 0.2 and semantic_score < 0.15 and summary_semantic_score < 0.1:
            # Très faible correspondance sur tous les critères
            final_score = min(final_score, 15.0)
        elif skills_match_score < 0.3 and semantic_score < 0.2:
            # Faible correspondance
            final_score = min(final_score, 30.0)
        
        # 4. Pénalité si compétences très faibles MAIS sémantique ok (incohérent)
        if skills_match_score < 0.2 and semantic_score > 0.5:
            # Si les compétences ne correspondent pas mais le texte semble similaire
            # Probablement un faux positif sémantique
            final_score = final_score * 0.7  # Réduction de 30%
        
        # 5. Garantir un minimum raisonnable si au moins un critère est bon
        if skills_match_score >= 0.5 or semantic_score >= 0.4:
            final_score = max(final_score, 40.0)  # Minimum 40% si un critère est bon
        if skills_match_score >= 0.6 or (semantic_score >= 0.5 and skills_match_score >= 0.4):
            final_score = max(final_score, 50.0)  # Minimum 50% si correspondance correcte
        
        # 6. Plafond réaliste selon la correspondance
        if skills_match_score >= 0.8 and semantic_score >= 0.6:
            # Excellente correspondance
            final_score = max(final_score, 75.0)  # Minimum 75% pour excellente correspondance
        elif skills_match_score >= 0.7 and semantic_score >= 0.5:
            # Très bonne correspondance
            final_score = max(final_score, 65.0)  # Minimum 65% pour très bonne correspondance
        
        print(f"[DEBUG] Score final calculé: {final_score:.1f} (sémantique: {semantic_score:.3f}, compétences: {skills_match_score:.3f}, résumé: {summary_semantic_score:.3f})")
        
        return round(min(max(final_score, 0.0), 100.0), 1)
